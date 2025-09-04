# doc_converters.py
import os, csv, json, zipfile, tempfile # Import tempfile here as well
from docx import Document
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import A4
from reportlab.lib.utils import simpleSplit
import pandas as pd
from openpyxl import Workbook
from PIL import Image, ImageDraw
from reportlab.platypus import SimpleDocTemplate, Paragraph, Preformatted, Table, TableStyle, Image as RLImage
from reportlab.platypus import Spacer
from pdf2image import convert_from_path # Corrected import
from PIL import Image
import pytesseract
from PyPDF2 import PdfMerger # Removed convert_from_path from this import

from reportlab.lib import colors
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.platypus import Paragraph
from reportlab.lib.units import cm
from reportlab.lib.utils import simpleSplit
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.pdfbase import pdfmetrics
from docx import Document as DocxReader
from docx.shared import Inches
import unicodedata # Import unicodedata for script detection
import shutil
import subprocess
from typing import List, Optional
import platform
import pathlib
import stat, requests, glob
import mammoth
from pdf2image import convert_from_path
import pdfkit
import json, re 
import os
import threading

# ---------- Page setup ----------
PAGE_W, PAGE_H = A4

# ---------------- Delete after 5 minutes ---------------- #
def schedule_delete(file_path, delay=300):  # 300 sec = 5 minutes
    def delete_file():
        if os.path.exists(file_path):
            try:
                if os.path.isdir(file_path):
                    import shutil
                    shutil.rmtree(file_path)
                else:
                    os.remove(file_path)
                print(f"🗑 Deleted: {file_path}")
            except Exception as e:
                print(f"❌ Failed to delete {file_path}: {e}")
    timer = threading.Timer(delay, delete_file)
    timer.start()

def doc_to_txt(docx_path, txt_path):
    try:
        doc = Document(docx_path)
        with open(txt_path, "w", encoding="utf-8") as out:
            # Paragraphs
            for para in doc.paragraphs:
                out.write((para.text or "") + "\n")
            # Tables (append as TSV)
            for table in doc.tables:
                for row in table.rows:
                    cells = [c.text.replace("\n"," ").strip() for c in row.cells]
                    out.write("\t".join(cells) + "\n")
        return txt_path
    except Exception as e:
      raise RuntimeError(f"Error converting DOCX to TXT: {e}") from e

# -------------------------
# Utilities
# -------------------------
def _which(cmd: str) -> Optional[str]:
    return shutil.which(cmd)

def _ensure_dir(path: str):
    os.makedirs(os.path.dirname(os.path.abspath(path)), exist_ok=True)

# -------------------------
# Script detection (used for font suggestions)
# -------------------------
def detect_script_simple(text: str) -> str:
    if not text:
        return "LATIN"
    for ch in text:
        cp = ord(ch)
        if 0x0900 <= cp <= 0x097F:  # Devanagari (Hindi, etc)
            return "DEVANAGARI"
        if (0x4E00 <= cp <= 0x9FFF) or (0x3400 <= cp <= 0x4DBF) or (0x3040 <= cp <= 0x30FF):
            return "CJK"
        if 0xAC00 <= cp <= 0xD7AF:
            return "HANGUL"
        if 0x0600 <= cp <= 0x06FF or 0x0750 <= cp <= 0x077F:
            return "ARABIC"
        if 0x0370 <= cp <= 0x03FF:
            return "GREEK"
    if any(ord(ch) > 127 for ch in text):
        return "OTHER"
    return "LATIN"

# ---------------- Helpers: Fonts -----------------
FONTS_DIR = "fonts"

# -------------------------
# Convert .doc -> .docx
# -------------------------
def convert_doc_to_docx_if_needed(input_path: str) -> str:
    """
    Converts .doc to .docx using LibreOffice (headless). Returns a path to a .docx file.
    Raises RuntimeError if soffice is missing or conversion fails.
    """
    if input_path.lower().endswith(".docx"):
        return input_path

    if not input_path.lower().endswith(".doc"):
        raise RuntimeError("Unsupported file (expecting .doc/.docx): " + input_path)

    soffice = _which("soffice")
    if not soffice:
        raise RuntimeError(
            "LibreOffice (soffice) not found. Install it to support .doc -> .docx conversion.\n"
            "Linux/Colab:  sudo apt-get update && sudo apt-get install -y libreoffice"
        )

    tmp_out = tempfile.mkdtemp(prefix="doc2docx_")
    cmd = [soffice, "--headless", "--convert-to", "docx", "--outdir", tmp_out, input_path]
    subprocess.check_call(cmd)
    out = os.path.join(tmp_out, os.path.splitext(os.path.basename(input_path))[0] + ".docx")
    if not os.path.exists(out):
        raise RuntimeError("LibreOffice did not produce DOCX.")
    return out

# -------------------------
# Attempt to install Noto fonts (best-effort for Linux/Colab)
# -------------------------
def ensure_noto_fonts_for_scripts(scripts: set):
    """
    Best-effort: download & install a couple of Noto fonts into user fonts dir on Linux.
    Only does anything on Linux-like systems (including Colab).
    For CJK (Chinese/Japanese/Korean) we print instructions because those packages are large.
    """
    if not scripts:
        return

    sysplat = platform.system().lower()
    if sysplat not in ("linux", "darwin"):
        # macOS/Windows: can't auto install reliably; give instructions
        print("Note: automatic font installation only supported for Linux/Colab in this script.")
        print("If you need CJK or other fonts installed on your system, please install Noto fonts manually.")
        return

    fonts_dir = os.path.expanduser("~/.local/share/fonts")
    os.makedirs(fonts_dir, exist_ok=True)

    # URLs for fonts (GoogleFonts repo raw)
    FONT_URLS = {
        "CJK":        "https://github.com/googlefonts/noto-cjk/raw/main/Sans/OTF/SimplifiedChinese/NotoSansSC-Regular.otf",
        "HANGUL":     "https://github.com/googlefonts/noto-cjk/raw/main/Sans/OTF/Korean/NotoSansKR-Regular.otf",
        "ARABIC":     "https://github.com/googlefonts/noto-fonts/raw/main/hinted/ttf/NotoNaskhArabic/NotoNaskhArabic-Regular.ttf",
        "GREEK":      "https://github.com/googlefonts/noto-fonts/raw/main/hinted/ttf/NotoSans/NotoSans-Regular.ttf",
        "OTHER":      "https://github.com/googlefonts/noto-fonts/raw/main/hinted/ttf/NotoSans/NotoSans-Regular.ttf",
        "LATIN": "https://github.com/googlefonts/noto-fonts/raw/main/hinted/ttf/NotoSans/NotoSans-Regular.ttf",
        "DEVANAGARI": "https://github.com/googlefonts/noto-fonts/raw/main/hinted/ttf/NotoSansDevanagari/NotoSansDevanagari-Regular.ttf",
        # CJK heavy: recommend apt-get fonts-noto-cjk instead of downloading single files
    }

    downloaded = []
    for s in scripts:
        if s == "CJK":
            # prefer apt package
            print("CJK script detected. For best results install Noto CJK fonts on your system.")
            print("On Debian/Ubuntu you can run:")
            print("  sudo apt-get update && sudo apt-get install -y fonts-noto-cjk")
            continue
        url = FONT_URLS.get(s, FONT_URLS.get("LATIN"))
        if not url:
            continue
        try:
            fname = os.path.basename(url)
            out_path = os.path.join(fonts_dir, fname)
            if os.path.exists(out_path):
                print(f"Font already present: {out_path}")
                downloaded.append(out_path)
                continue
            import urllib.request
            print("Downloading font:", url)
            urllib.request.urlretrieve(url, out_path)
            os.chmod(out_path, stat.S_IRUSR | stat.S_IWUSR | stat.S_IRGRP | stat.S_IROTH)
            downloaded.append(out_path)
            print(" → saved to", out_path)
        except Exception as e:
            print("Failed to download font:", e)

    # refresh font cache (Linux)
    try:
        print("Refreshing font cache...")
        run(["fc-cache", "-f", "-v"])
    except Exception as e:
        print("fc-cache failed (non-fatal):", e)

    if downloaded:
        print("Fonts installed (user-level). Restart your PDF viewer if needed.")

# -------------------------
# Convert .docx -> .pdf using LibreOffice (preferred)
# -------------------------
def convert_docx_to_pdf_libreoffice(docx_path: str, output_pdf: str) -> str:
    soffice = _which("soffice")
    if not soffice:
        raise RuntimeError(
            "LibreOffice (soffice) not found. Install it for high-fidelity DOCX->PDF.\n"
            "Linux/Colab:  sudo apt-get update && sudo apt-get install -y libreoffice"
        )
    outdir = os.path.dirname(os.path.abspath(output_pdf)) or "."
    _ensure_dir(output_pdf)
    cmd = [soffice, "--headless", "--convert-to", "pdf", "--outdir", outdir, docx_path]
    subprocess.check_call(cmd)
    produced = os.path.join(outdir, os.path.splitext(os.path.basename(docx_path))[0] + ".pdf")
    if not os.path.exists(produced):
        raise RuntimeError("LibreOffice did not produce PDF.")
    if os.path.abspath(produced) != os.path.abspath(output_pdf):
        os.replace(produced, output_pdf)
    return output_pdf

# ---------- Rasterize (guaranteed readability) ----------
def rasterize_pdf_to_pdf(input_pdf: str, output_pdf: str, dpi: int = 200) -> str:
    """
    Converts every page of PDF to an image, then merges them back into PDF.
    Ensures fonts, multilingual text, tables, JSON, etc. are preserved visually.
    """
    _ensure_dir(output_pdf)
    images = convert_from_path(input_pdf, dpi=dpi)
    if not images:
        raise RuntimeError("No pages produced while rasterizing.")

    pdf_pages = []
    for img in images:
        if img.mode != "RGB":
            img = img.convert("RGB")
        pdf_pages.append(img)

    first, rest = pdf_pages[0], pdf_pages[1:]
    first.save(output_pdf, save_all=True, append_images=rest)
    return output_pdf


# ---------- Main unified function ----------
def doc_to_pdf(input_path: str,
               output_pdf: str,
               mode: str = "auto",
               raster_dpi: int = 200) -> str:
    """
    Convert .doc/.docx → .pdf

    - mode="auto":
        * If doc has table(s) with >=5 columns → convert via HTML to PDF (avoids cropping).
        * Else → direct via LibreOffice (fastest, keeps structure).
    - mode="image": LibreOffice → PDF → rasterized (pixel perfect, no font issues).
    """
    if not os.path.exists(input_path):
        raise RuntimeError("Input file not found: " + input_path)

    _ensure_dir(output_pdf)

    # 1) Ensure DOCX
    docx_path = convert_doc_to_docx_if_needed(input_path)

    # --- Check if doc has wide tables (>=5 columns) ---
    has_wide_tables = False
    try:
        from docx import Document
        doc = Document(docx_path)
        for tbl in doc.tables:
            if tbl.rows and len(tbl.rows[0].cells) >= 5:
                has_wide_tables = True
                break
    except Exception as e:
        print(f"⚠️ Could not inspect tables in DOCX: {e}")

    # --- Wide tables? Use HTML → PDF ---
    if has_wide_tables and mode == "auto":
        try:

            with open(docx_path, "rb") as f:
                html = mammoth.convert_to_html(f).value
            tmp_html = os.path.join(tempfile.gettempdir(), "doc_tmp.html")
            with open(tmp_html, "w", encoding="utf-8") as f:
                f.write(html)
            pdfkit.from_file(tmp_html, output_pdf)
            print(f"✅ Converted via HTML path (wide table support): {output_pdf}")
            return output_pdf
        except Exception as e:
            print(f"⚠️ HTML conversion failed, falling back to LibreOffice: {e}")

    # --- Default: DOCX → PDF via LibreOffice ---
    tmp_dir = tempfile.mkdtemp(prefix="docx2pdf_")
    lo_pdf = os.path.join(tmp_dir, "lo.pdf")
    try:
        lo_pdf = convert_docx_to_pdf_libreoffice(docx_path, lo_pdf)
    except Exception as e:
        shutil.rmtree(tmp_dir)
        raise RuntimeError(f"LibreOffice conversion failed: {e}")

    if mode == "auto":
        shutil.copy(lo_pdf, output_pdf)
        shutil.rmtree(tmp_dir)
        return output_pdf

    if mode == "image":
        raster_pdf = os.path.join(tmp_dir, "raster.pdf")
        rasterize_pdf_to_pdf(lo_pdf, raster_pdf, dpi=raster_dpi)
        shutil.copy(raster_pdf, output_pdf)
        shutil.rmtree(tmp_dir)
        return output_pdf

    shutil.rmtree(tmp_dir)
    raise ValueError('mode must be "auto" or "image"')


#--------- Docx For Image---------------------
def doc_to_docx_image(input_path: str) -> str:
    """
    Ensure the input is in DOCX format.
    If it's a .doc, convert using LibreOffice headless.
    """
    import tempfile, subprocess, os, glob

    if input_path.lower().endswith(".docx"):
        return input_path

    if input_path.lower().endswith(".doc"):
        tmp_dir = tempfile.mkdtemp(prefix="doc2docx_")
        try:
            subprocess.check_call([
                "soffice", "--headless", "--convert-to", "docx",
                "--outdir", tmp_dir, input_path
            ])

            # ✅ Find the .docx file inside tmp_dir
            converted_files = glob.glob(os.path.join(tmp_dir, "*.docx"))
            if not converted_files:
                raise RuntimeError("DOC→DOCX conversion failed, file not created.")

            return converted_files[0]  # Return the actual converted path
        except Exception as e:
            raise RuntimeError(f"Failed to convert .doc to .docx: {e}")
    else:
        raise RuntimeError(f"Unsupported file type: {input_path}")

def doc_to_image(input_path, out_path, dpi=200):
    """
    Convert .doc/.docx → PNG image(s) with full formatting (tables, images, layout preserved).
    - If 1 page → single PNG
    - If multiple pages → user chooses:
        1 = merge into single long PNG
        2 = separate PNG per page (zipped)
    """
    tmp_dir = None
    try:
        # ✅ Always ensure DOCX
        docx_path = doc_to_docx_image(input_path)
        # ✅ Detect wide tables
        wide_table = False
        try:
            doc = Document(docx_path)
            for table in doc.tables:
                if len(table.columns) > 5:
                    wide_table = True
                    break
        except Exception as e:
            print(f"⚠️ Could not inspect tables: {e}")

        # ========== PATH 1: Wide Table → Mammoth + WeasyPrint ==========
        if wide_table:
            print("⚠️ Wide table detected → Using HTML → PDF → Image pipeline")
            tmp_dir = tempfile.mkdtemp(prefix="doc2html_")

            # DOCX → HTML
            with open(docx_path, "rb") as f:
                result = mammoth.convert_to_html(f)
                html_content = result.value

            # Save HTML to file (safe for large tables)
            html_file = os.path.join(tmp_dir, "doc.html")
            with open(html_file, "w", encoding="utf-8") as f:
                f.write(html_content)

            # HTML → PDF
            pdf_path = os.path.join(tmp_dir, "out.pdf")
            pdfkit.from_file(html_file, pdf_path)

        # ========== PATH 2: Normal → LibreOffice ==========
        else:
            print("✅ No wide tables → Using LibreOffice pipeline")
            tmp_dir = tempfile.mkdtemp(prefix="doc2pdf_")
            subprocess.check_call([
                "soffice", "--headless", "--convert-to", "pdf",
                "--outdir", tmp_dir, docx_path
            ])

            pdf_files = glob.glob(os.path.join(tmp_dir, "*.pdf"))
            if not pdf_files:
                raise RuntimeError("DOCX→PDF failed, no PDF generated.")
            pdf_path = pdf_files[0]

        print(f"✅ PDF ready at: {pdf_path}", '3')

        # ✅ PDF → PNG(s)
        pages = convert_from_path(pdf_path, dpi=dpi)
        print(f"pages: {len(pages)}", "4")

        if len(pages) == 1:
            if not out_path.lower().endswith(".png"):
                out_path += ".png"
            pages[0].save(out_path, "PNG")
            print(f"✅ Converted single page image: {out_path}")
            return out_path

        # Multiple pages → ask user
        print("Document has multiple pages. Choose output format:")
        print("1. Single PNG (all pages merged vertically)")
        print("2. Separate PNG per page (zipped)")
        choice = input("Enter choice (1/2): ").strip()

        if choice == "1":
            widths, heights = zip(*(p.size for p in pages))
            total_height = sum(heights)
            max_width = max(widths)
            merged_img = Image.new("RGB", (max_width, total_height), "white")
            y = 0
            for p in pages:
                merged_img.paste(p, (0, y))
                y += p.height
            if not out_path.lower().endswith(".png"):
                out_path += ".png"
            merged_img.save(out_path, "PNG")
            print(f"✅ Merged multi-page image saved: {out_path}")
            return out_path

        elif choice == "2":
            if not out_path.lower().endswith(".zip"):
                out_path += ".zip"
            img_dir = tempfile.mkdtemp(prefix="doc2img_")
            image_files = []
            for i, page in enumerate(pages, 1):
                img_path = os.path.join(img_dir, f"page_{i}.png")
                page.save(img_path, "PNG")
                image_files.append(img_path)

            with zipfile.ZipFile(out_path, "w") as zipf:
                for f in image_files:
                    zipf.write(f, os.path.basename(f))
            print(f"✅ Separate page images saved as zip: {out_path}")
            return out_path

        else:
            raise ValueError("Invalid choice")

    except Exception as e:
        raise RuntimeError(f"Error converting DOCX to PNG: {e}") from e
    finally:
        if tmp_dir and os.path.exists(tmp_dir):
            shutil.rmtree(tmp_dir, ignore_errors=True)

def doc_to_csv(input_path, csv_path):
    try:
        # Step 1: normalize DOC → DOCX
        docx_path = doc_to_docx_image(input_path)
        doc = Document(docx_path)

        # 🚫 Step 2: Reject if any images exist
        for rel in doc.part.rels.values():
            if "image" in rel.target_ref:
                raise RuntimeError("❌ Not supported: Document contains images, cannot export to CSV.")

        # Step 3: Scan all tables first
        valid_tables = []
        for table in doc.tables:
            col_count = len(table.rows[0].cells) if table.rows else 0
            row_count = len(table.rows)

            if col_count < 4 or row_count < 2:
                continue  

            meaningful_rows = 0
            for row in table.rows:
                texts = [c.text.strip() for c in row.cells if c.text.strip()]
                if len(texts) >= (col_count // 2) and any(len(t) >= 3 for t in texts):
                    meaningful_rows += 1

            if meaningful_rows >= 2:
                valid_tables.append(table)

        if not valid_tables:
            raise RuntimeError("❌ Not supported: No valid table (≥4 columns, ≥2 rows, with real data).")

        # ✅ Step 4: Now open CSV only if conversion is confirmed
        with open(csv_path, "w", encoding="utf-8-sig", newline="") as out:
            writer = csv.writer(out)
            for table in valid_tables:
                for row in table.rows:
                    row_data = []
                    for c in row.cells:
                        text = c.text.replace("\n", " ").strip()
                        if text:
                            text = text.encode("utf-8", "ignore").decode("utf-8")
                        row_data.append(text)
                    writer.writerow(row_data)

        return csv_path

    except Exception as e:
        # ❌ Ensure partial/empty CSV is deleted if error
        if os.path.exists(csv_path):
            os.remove(csv_path)
        raise RuntimeError(f"Error converting DOC/DOCX to CSV: {e}") from e

def doc_to_xls(input_path, xls_path):
    try:
        # ✅ DOC → DOCX ensure
        docx_path = doc_to_docx_image(input_path)

        # ✅ Load DOCX
        doc = Document(docx_path)

        # ✅ Image check (unsupported for XLS export)
        for shape in doc.inline_shapes:
            raise RuntimeError("❌ Not supported: Document contains images, cannot convert to XLS.")

        wrote_any = False
        valid_table_found = False

        wb = Workbook(write_only=True)
        ws = wb.create_sheet("Sheet1")

        # ✅ Table scan
        for table in doc.tables:
            if not table.rows or len(table.rows[0].cells) < 4:
                continue  # skip invalid
            if len(table.rows) < 2:
                continue  # need at least 2 rows

            # check real data
            has_data = any(any(c.text.strip() for c in row.cells) for row in table.rows)
            if not has_data:
                continue

            valid_table_found = True
            for row in table.rows:
                row_data = []
                for c in row.cells:
                    text = c.text.replace("\n", " ").strip()
                    if text:
                        text = text.encode("utf-8", "ignore").decode("utf-8")
                    row_data.append(text)
                ws.append(row_data)
                wrote_any = True

        # ❌ No valid table
        if not valid_table_found:
            raise RuntimeError("❌ Not supported: No valid table (≥4 columns, ≥2 rows, with real data).")

        if not wrote_any:
            raise RuntimeError("❌ Table found but empty, cannot convert.")

        wb.save(xls_path)
        return xls_path

    except Exception as e:
        # ⚠️ Agar error ho to partial file delete
        if os.path.exists(xls_path):
            os.remove(xls_path)
        raise RuntimeError(f"Error converting DOC/DOCX to XLSX: {e}") from e

def doc_to_json(input_path, json_path):
    try:
        # ✅ DOC → DOCX ensure
        docx_path = doc_to_docx_image(input_path)

        # ✅ Load DOCX
        doc = Document(docx_path)

        # ✅ Step 3: Image check
        for rel in doc.part.rels.values():
            if "image" in rel.reltype:
                raise RuntimeError("❌ Not supported: Document contains images.")

        # ✅ Step 4: Raw text join
        raw_text = "\n".join([p.text.strip() for p in doc.paragraphs if p.text.strip()])

        # ✅ Step 5: Detect JSON-like content
        if raw_text.strip().startswith(("{", "[")):
            try:
                parsed = json.loads(raw_text)
                with open(json_path, "w", encoding="utf-8") as out:
                    json.dump(parsed, out, ensure_ascii=False, indent=2)
                return json_path
            except Exception:
                pass  # not valid JSON, fallback

        # ✅ Step 6: Detect table data
        valid_tables = []
        for table in doc.tables:
            if len(table.columns) >= 4 and len(table.rows) >= 2:
                table_data = []
                for row in table.rows:
                    row_data = []
                    for c in row.cells:
                        text = c.text.replace("\n", " ").strip()
                        if text:
                            text = text.encode("utf-8", "ignore").decode("utf-8")
                        row_data.append(text)
                    table_data.append(row_data)
                valid_tables.append(table_data)

        # ✅ Step 7: Detect tab/comma separated plain text (like csv.docx)
        if not valid_tables:
            lines = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
            table_data = []
            for line in lines:
                if "," in line or "\t" in line or ";" in line:
                    parts = re.split(r'[\t,;]', line)
                    if len(parts) >= 4:
                        table_data.append([p.strip() for p in parts])
            if table_data:
                valid_tables.append(table_data)

        # ✅ Step 8: Final check
        if not valid_tables:
            raise RuntimeError("❌ Not supported: No valid JSON or tabular data found.")

        payload = {"tables": valid_tables}
        with open(json_path, "w", encoding="utf-8") as out:
            json.dump(payload, out, ensure_ascii=False, indent=2)

        return json_path


    except Exception as e:
        # ⚠️ Error aane par partial file delete
        if os.path.exists(json_path):
            os.remove(json_path)
        raise RuntimeError(f"Error converting DOC/DOCX to JSON: {e}") from e
        
# ---------------- CLI MENU ---------------- #

def main():
    # For Colab file handling
    try:
        from google.colab import files
    except ImportError:
        files = None

    while True:
        print("\n==== DOCX File Converter =====")
        print("1. Convert DOCX File")
        print("2. Exit")
        choice = input("Enter choice: ").strip()

        if choice == "1":
            # Upload file (Colab) or enter path (offline)
            if files:
                uploaded = files.upload()
                pdf_file = list(uploaded.keys())[0]
                #pdf_file = input("Enter path of DOCX/DOC file: ").strip()
            else:
                pdf_file = input("Enter path of DOC/DOCX file: ").strip()

            print("\nSupported conversions: ")
            print("1. DOC → PDF")
            print("2. DOC → TXT")
            print("3. DOC → PNG (image)")
            print("4. DOC → CSV")
            print("5. DOC → XLSX")
            print("6. DOC → JSON")
            fmt_choice = input("Select target format (1-6): ").strip()

            base, _ = os.path.splitext(pdf_file)
            out_file = None

            try:
                if fmt_choice == "1":
                    out_file = base + ".pdf"
                    doc_to_pdf(pdf_file, out_file)
                elif fmt_choice == "2":
                    out_file = base + ".txt"
                    doc_to_txt(pdf_file, out_file)
                elif fmt_choice == "3":
                    out_file = base + ".png"
                    out_dir = base + "_images"
                    # Modified to use the corrected doc_to_image
                    out_file = doc_to_image(pdf_file, out_file)
                elif fmt_choice == "4":
                    out_file = base + ".csv"
                    doc_to_csv(pdf_file, out_file)
                elif fmt_choice == "5":
                    out_file = base + ".xlsx"
                    doc_to_xls(pdf_file, out_file)
                elif fmt_choice == "6":
                    out_file = base + ".json"
                    doc_to_json(pdf_file, out_file)
                else:
                    print("❌ Invalid choice!")
                    continue

                print(f"✅ Converted successfully: {out_file}")

                # Schedule auto-delete after 5 minutes
                if out_file:
                    schedule_delete(out_file)

                # Colab download option
                if files and out_file:
                    if os.path.isdir(out_file): # if output is a directory (images)
                         print(f"💡 Multiple files saved in {out_file}. You may need to zip and download manually.")
                         # Example of zipping and offering download (requires zip installed)
                         # !zip -r {out_file}.zip {out_file}
                         # files.download(f"{out_file}.zip")
                    else:
                         files.download(out_file)


            except Exception as e:
                print("❌ Conversion failed:", e)

        elif choice == "2":
            print("👋 Exiting...")
            break
        else:
            print("❌ Invalid choice!")

if __name__ == "__main__":
    main()