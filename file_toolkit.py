#!/usr/bin/env python3
# unified_converter.py

import os
import sys
import csv
import json
import argparse
import threading
import textwrap
import zipfile
import shutil
from io import StringIO
import math
import io
import unicodedata
import uuid
import ijson
import unicodedata # Import unicodedata for script detection
import shutil
import subprocess
import platform
import pathlib
import stat, requests, glob

# Third-party libs
import langid
import pytesseract
import pdfplumber
import unicodedata
import pandas as pd
import pdfkit
import pypandoc
import mammoth
from tqdm import tqdm
import langid
from weasyprint import HTML
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
import easyocr # Import easyocr
from typing import List, Union, Optional
from langdetect import detect
from docx.shared import Inches
from docx.oxml import OxmlElement
from openpyxl.styles import Font
from reportlab.pdfgen import canvas
from reportlab.lib import colors
from reportlab.lib.units import cm
from pdf2image import convert_from_path
from reportlab.platypus import Spacer
from reportlab.pdfbase import pdfmetrics
from deep_translator import GoogleTranslator
from reportlab.lib.utils import simpleSplit
from openpyxl import Workbook, load_workbook
from docx import Document as DocxReader
from reportlab.platypus import Paragraph
from PIL import Image, ImageEnhance, ImageFilter
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.lib.pagesizes import A4, letter
from reportlab.pdfgen import canvas as rl_canvas
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.pdfbase.cidfonts import UnicodeCIDFont
from pdfminer.high_level import extract_text as pdf_extract_text
from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Preformatted, Image as RLImage
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from langdetect import detect
import regex  # better than re for Unicode script detection
import easyocr # Import easyocr
from PyPDF2 import PdfMerger # Removed convert_from_path from this import



PAGE_W, PAGE_H = A4

# =========================
# Utilities
# =========================

def _which(cmd: str) -> Optional[str]:
    return shutil.which(cmd)

# Optional libs
try:
    from langdetect import detect, detect_langs
    _HAS_LANGDETECT = True
except Exception:
    _HAS_LANGDETECT = False

# pandas only if user wants xlsx export or fallback viewing
try:
    import pandas as pd
    _HAS_PANDAS = True
except Exception:
    _HAS_PANDAS = False

try:
    import tkinter as tk
    from tkinter import messagebox
    _HAS_TK = True
except Exception:
    _HAS_TK = False

def _auto_detect_delimiter(txt_path, sample_bytes=8192):
    """Try to detect delimiter from first KBs of file; fallback to tab."""
    try:
        with open(txt_path, "r", encoding="utf-8", errors="ignore") as f:
            sample = f.read(sample_bytes)
        sniffer = csv.Sniffer()
        dialect = sniffer.sniff(sample)
        return dialect.delimiter
    except Exception:
        # Common fallbacks
        for d in [",", "\t", ";", "|"]:
            if d in sample:
                return d
    return "\t"

def _detect_non_english_and_scripts(txt_path, max_chars=10000):
    """
    Returns tuple (is_non_english_bool, top_languages_list, has_non_ascii_bool, scripts_set)
    Uses langdetect if available; always returns heuristic based on non-ASCII presence.
    """
    text_sample = ""
    with open(txt_path, "r", encoding="utf-8", errors="ignore") as f:
        # read some lines to keep it fast for very large files
        for i, line in enumerate(f):
            text_sample += line
            if len(text_sample) >= max_chars:
                break

    has_non_ascii = any(ord(ch) > 127 for ch in text_sample)
    top_langs = []
    non_english = False

    if _HAS_LANGDETECT:
        try:
            langs = detect_langs(text_sample)
            top_langs = [(str(l).split(":")[0], float(str(l).split(":")[1])) for l in langs]
            # consider non-english if top language not 'en' with decent prob
            if top_langs:
                top_lang = top_langs[0][0]
                top_prob = top_langs[0][1]
                non_english = not (top_lang == "en" and top_prob >= 0.65)
        except Exception:
            top_langs = []
            non_english = has_non_ascii
    else:
        non_english = has_non_ascii

    # simple script heuristics (return set like {'Devanagari','CJK','Arabic',...})
    scripts = set()
    for ch in text_sample[:2000]:
        cp = ord(ch)
        if 0x0900 <= cp <= 0x097F:
            scripts.add("Devanagari")
        elif 0x4E00 <= cp <= 0x9FFF:
            scripts.add("CJK")
        elif 0x3040 <= cp <= 0x30FF:
            scripts.add("Japanese")
        elif 0xAC00 <= cp <= 0xD7AF:
            scripts.add("Hangul")
        elif 0x0600 <= cp <= 0x06FF:
            scripts.add("Arabic")
        elif cp > 127 and not ch.isspace():
            scripts.add("OtherNonASCII")

    return non_english, top_langs, has_non_ascii, scripts

def _notify_user_excel_instructions(csv_path, non_english, top_langs, scripts):
    """
    Shows instructions to user on how to open CSV in Excel correctly.
    If running in Colab -> just prints instructions. If local & tkinter available -> show popup.
    """
    languages_text = ""
    if top_langs:
        languages_text = ", ".join([f"{l}:{p:.2f}" for l, p in top_langs])
    else:
        languages_text = "detected non-ASCII text" if non_english else "English/ASCII"

    instructions = (
        f"CSV saved at: {csv_path}\n\n"
        f"Detected languages/sample: {languages_text}\n\n"
        "For best results open this CSV in Excel like this:\n"
        "1) Open Excel -> Data tab -> Get Data -> From Text/CSV\n"
        "2) Select the file and when the preview appears, choose 'File origin' = UTF-8 (or '65001: Unicode (UTF-8)')\n"
        "3) Click 'Load'.\n\n"
        "Alternative (no import wizard): this script saved the file with a BOM (utf-8-sig) so modern Excel often opens it fine by double-clicking.\n\n"
        "If users still see garbled characters, consider opening in LibreOffice or importing via the Data menu and explicitly choosing UTF-8.\n\n"
        "Tip: To avoid Excel problems at all, open the provided .xlsx version (if available) which preserves texts in all languages."
    )

    # If Colab environment, just print
    try:
        import google.colab  # type: ignore
        in_colab = True
    except Exception:
        in_colab = False

    if in_colab or not _HAS_TK:
        print("\n" + "="*60 + "\nIMPORTANT: How to open this CSV in Excel\n" + "="*60)
        print(instructions)
        return

    # else show a tkinter popup (local desktop)
    try:
        root = tk.Tk()
        root.withdraw()
        # Use a scrolled text popup? Simpler: showinfo with basic instructions
        messagebox.showinfo("CSV saved — How to open in Excel (UTF-8)", instructions)
        root.destroy()
    except Exception:
        print(instructions)

def schedule_delete(file_path, delay=300):
    """Auto-delete a file/folder after delay seconds."""
    try:
        def delete_file():
            if os.path.exists(file_path):
                try:
                    if os.path.isdir(file_path):
                        shutil.rmtree(file_path)
                    else:
                        os.remove(file_path)
                    print(f"🗑 Deleted: {file_path}")
                except Exception as e:
                    print(f"❌ Failed to delete {file_path}: {e}")
        timer = threading.Timer(delay, delete_file)
        timer.start()
    except Exception as e:
        print(f"⚠️ schedule_delete failed: {e}")

def ensure_parent_dir(path: str):
    try:
        parent = os.path.dirname(os.path.abspath(path))
        if parent and not os.path.exists(parent):
            os.makedirs(parent, exist_ok=True)
    except Exception as e:
        print(f"⚠️ Could not ensure parent dir for {path}: {e}")

# ---------- Script to Font mapping (for image conversions) ----------
FALLBACK_FONTS = {
    "LATIN": "arial.ttf",
    "DEVANAGARI": "NotoSansDevanagari-Medium.ttf",
    "CJK": "arialuni.ttf",
    "ARABIC": "arial.ttf",
    "GREEK": "arial.ttf",
    "OTHER": "arialuni.ttf",
    "DEFAULT": "arial.ttf"
}

def detect_script_text(text):
    """Detect script of a line and return font name."""
    # Register multiple fonts
    pdfmetrics.registerFont(TTFont("NotoSans", "/content/NotoSans-Regular.ttf")) #Add Local path of font
    pdfmetrics.registerFont(TTFont("NotoSansDevanagari", "/content/NotoSansDevanagari-Medium.ttf")) #Add Local path of font

    if regex.search(r'\p{Devanagari}', text):
        return "NotoSansDevanagari"
    return "NotoSans"  # default English font

def get_font_for_line(line, font_size):
    # font mapping by language
    FONT_MAP = {
        "hi": "/content/NotoSansDevanagari-Medium.ttf",  # Hindi
        "en": "/content/NotoSans-Regular.ttf",           # English
        "zh-cn": "/content/NotoSansSC-Regular.ttf",      # Simplified Chinese
        "ja": "/content/NotoSansJP-Regular.ttf",         # Japanese
        "ko": "/content/NotoSansKR-Regular.ttf",         # Korean
        "default": "/content/NotoSans-Regular.ttf"       # fallback
    }

    try:
        lang = detect(line) if line.strip() else "en"
    except:
        lang = "en"

    # normalize
    if lang.startswith("zh"):
        lang = "zh-cn"
    elif lang.startswith("ja"):
        lang = "ja"
    elif lang.startswith("ko"):
        lang = "ko"
    elif lang.startswith("hi"):
        lang = "hi"
    else:
        lang = "en"

    font_path = FONT_MAP.get(lang, FONT_MAP["default"])
    if not os.path.exists(font_path):  # fallback if missing
        font_path = FONT_MAP["default"]

    return ImageFont.truetype(font_path, font_size)

def flatten_json(obj, parent_key="", sep="."):
    """Flatten nested JSON (dicts/lists) into key-value pairs."""
    items = []
    if isinstance(obj, dict):
        for k, v in obj.items():
            new_key = f"{parent_key}{sep}{k}" if parent_key else k
            items.extend(flatten_json(v, new_key, sep=sep).items())
    elif isinstance(obj, list):
        for i, v in enumerate(obj):
            new_key = f"{parent_key}[{i}]"
            items.extend(flatten_json(v, new_key, sep=sep).items())
    else:
        items.append((parent_key, obj))
    return dict(items)

def is_ndjson(json_path):
    """Detect NDJSON (multiple lines with valid JSON objects)"""
    with open(json_path, "r", encoding="utf-8") as f:
        try:
            json.load(f)   # normal JSON
            return False
        except json.JSONDecodeError:
            return True

def detect_script_simple(text: str) -> str:
    if not text:
        return "LATIN"
    for ch in text:
        cp = ord(ch)
        # Devanagari
        if 0x0900 <= cp <= 0x097F:
            return "DEVANAGARI"
        # CJK unified ideographs (Chinese)
        if 0x4E00 <= cp <= 0x9FFF or 0x3400 <= cp <= 0x4DBF:
            return "CJK"
        # Hiragana / Katakana / Kanji ranges also considered CJK/Japanese
        if 0x3040 <= cp <= 0x30FF:
            return "CJK"
        # Hangul (Korean)
        if 0xAC00 <= cp <= 0xD7AF:
            return "HANGUL"
        # Arabic
        if 0x0600 <= cp <= 0x06FF or 0x0750 <= cp <= 0x077F:
            return "ARABIC"
        # Greek
        if 0x0370 <= cp <= 0x03FF:
            return "GREEK"
        # Basic Latin (ASCII)
        if cp <= 0x007F:
            continue
        # If many other non-ascii chars, mark OTHER as fallback
    # fallback: if there are non-ascii chars but not matched above
    if any(ord(ch) > 127 for ch in text):
        return "OTHER"
    return "LATIN"

def get_font(script, default_size=12):
    font_name = FALLBACK_FONTS.get(script, FALLBACK_FONTS["DEFAULT"])
    search_paths = [
        f"/usr/share/fonts/truetype/liberation/{font_name}",
        f"/usr/share/fonts/truetype/msttcorefonts/{font_name}",
        f"/usr/share/fonts/truetype/noto/{font_name}",
        f"/Library/Fonts/{font_name}",
        f"C:/Windows/Fonts/{font_name}",
        font_name
    ]
    for p in search_paths:
        try:
            if os.path.exists(p):
                return ImageFont.truetype(p, default_size)
        except Exception:
            pass
    # Fallbacks
    for fb in ["arialuni.ttf", "NotoSans-Regular.ttf"]:
        for p in search_paths:
            try:
                alt = p.replace(font_name, fb)
                if os.path.exists(alt):
                    return ImageFont.truetype(alt, default_size)
            except Exception:
                pass
    # Last resort
    try:
        return ImageFont.load_default()
    except Exception as e:
        raise RuntimeError(f"Could not load any font: {e}")

def get_font_path(script, default_size=12):
    font_name = FALLBACK_FONTS.get(script, FALLBACK_FONTS["DEFAULT"])
    # In a Colab environment, commonly available fonts are in specific paths
    # You might need to adjust these paths or download fonts if needed.
    # Example paths (may vary):
    colab_font_paths = [
        f"/usr/share/fonts/truetype/liberation/{font_name}", # Liberation fonts often present
        f"/usr/share/fonts/truetype/msttcorefonts/{font_name}", # If ttf-mscorefonts-installer is run
        f"/usr/share/fonts/truetype/noto/{font_name}", # Noto fonts
        f"/content/{font_name}", # If user uploaded
        font_name # As a last resort, try just the name
    ]

    for font_path in colab_font_paths:
        if os.path.exists(font_path):
            try:
                return ImageFont.truetype(font_path, default_size)
            except Exception as e:
                print(f"⚠️ Could not load font {font_path}: {e}")

    # If specific font not found, try a generic fallback font that supports Unicode
    try:
        # Try a widely available Unicode font like Arial Unicode MS or Noto Sans
        fallback_unicode_fonts = ["arialuni.ttf", "NotoSans-Regular.ttf"]
        for fb_font_name in fallback_unicode_fonts:
            for fb_path in colab_font_paths:
                 potential_path = fb_path.replace(font_name, fb_font_name)
                 if os.path.exists(potential_path):
                     try:
                         return ImageFont.truetype(potential_path, default_size)
                     except Exception:
                         pass # Try next fallback
    except Exception:
        pass # Continue to reportlab fallback

    # Fallback to ReportLab's default font if Pillow fails
    print(f"⚠️ Font {font_name} not found or could not be loaded. Using default Pillow font.")
    try:
        # Attempt to return a basic Pillow font if no TTF found
        return ImageFont.load_default()
    except Exception as e:
        print(f"❌ Could not load default Pillow font: {e}")
        raise RuntimeError("Could not load any font.")


def wrap_text(line: str, font_size: int, img_width: int, margin: int) -> list:
    # empirical width per char estimate
    max_chars = max(1, int((img_width - 2 * margin) / (font_size * 0.6)))
    return textwrap.wrap(line, width=max_chars)

# ---- Script to Font mapping for PDF----
FALLBACK_FONTS_PDF = {
    "LATIN": "Times New Roman",
    "DEVANAGARI": "Nirmala UI",        # Hindi
    "CJK": "SimSun",                   # Chinese/Japanese/Korean
    "ARABIC": "Amiri",
    "GREEK": "Palatino Linotype",
    "DEFAULT": "Arial"
}

def _page_tables_to_rows(page):
    """
    Try to extract tables from a pdfplumber page as rows.
    If no tables detected, fall back to extracting text lines and splitting heuristically.
    Returns a list of row-lists.
    """
    rows = []
    # 1) Try extract_tables (pdfplumber)
    try:
        tables = page.extract_tables()  # returns list of tables, each table is list of rows
        if tables:
            for table in tables:
                # table: list of rows (cells may be None)
                for row in table:
                    rows.append([cell for cell in row])
            if rows:
                return rows
    except Exception:
        # continue to fallback
        pass

    # 2) Fallback: use page.extract_text() and split lines by common delimiters (tab or comma)
    text = page.extract_text() or ""
    if not text:
        return []  # nothing to do
    # attempt to detect delimiter by checking first non-empty line
    lines = [ln for ln in text.splitlines() if ln.strip()]
    if not lines:
        return []
    first = lines[0]
    # choose delimiter heuristically
    delim = "\t"
    if "," in first and first.count(",") >= 1:
        delim = ","
    elif "|" in first and first.count("|") >= 1:
        delim = "|"
    # split each line
    for ln in lines:
        parts = [p.strip() for p in ln.split(delim)]
        rows.append(parts)
    return rows

def detect_script_pdf(text):
    """Detects script of the text based on Unicode names"""
    for ch in text:
        name = unicodedata.name(ch, "")
        if "DEVANAGARI" in name:
            return "DEVANAGARI"
        elif any(x in name for x in ["CJK UNIFIED", "HIRAGANA", "KATAKANA", "HANGUL"]):
            return "CJK"
        elif "ARABIC" in name:
            return "ARABIC"
        elif "GREEK" in name:
            return "GREEK"
    return "LATIN"

# ---------------- Helpers: Fonts -----------------
FONTS_DIR = "fonts"

# -------------------------
# Convert .doc -> .docx
# -------------------------
def convert_doc_to_docx_if_needed(input_path: str) -> str:
    """
    Converts .doc to .docx using LibreOffice (headless). Returns a path to a .docx file.
    Raises RuntimeError if soffice is missing or conversion fails.
    """
    if input_path.lower().endswith(".docx"):
        return input_path

    if not input_path.lower().endswith(".doc"):
        raise RuntimeError("Unsupported file (expecting .doc/.docx): " + input_path)

    soffice = _which("soffice")
    if not soffice:
        raise RuntimeError(
            "LibreOffice (soffice) not found. Install it to support .doc -> .docx conversion.\n"
            "Linux/Colab:  sudo apt-get update && sudo apt-get install -y libreoffice"
        )

    tmp_out = tempfile.mkdtemp(prefix="doc2docx_")
    cmd = [soffice, "--headless", "--convert-to", "docx", "--outdir", tmp_out, input_path]
    subprocess.check_call(cmd)
    out = os.path.join(tmp_out, os.path.splitext(os.path.basename(input_path))[0] + ".docx")
    if not os.path.exists(out):
        raise RuntimeError("LibreOffice did not produce DOCX.")
    return out

# -------------------------
# Attempt to install Noto fonts (best-effort for Linux/Colab)
# -------------------------
def ensure_noto_fonts_for_scripts(scripts: set):
    """
    Best-effort: download & install a couple of Noto fonts into user fonts dir on Linux.
    Only does anything on Linux-like systems (including Colab).
    For CJK (Chinese/Japanese/Korean) we print instructions because those packages are large.
    """
    if not scripts:
        return

    sysplat = platform.system().lower()
    if sysplat not in ("linux", "darwin"):
        # macOS/Windows: can't auto install reliably; give instructions
        print("Note: automatic font installation only supported for Linux/Colab in this script.")
        print("If you need CJK or other fonts installed on your system, please install Noto fonts manually.")
        return

    fonts_dir = os.path.expanduser("~/.local/share/fonts")
    os.makedirs(fonts_dir, exist_ok=True)

    # URLs for fonts (GoogleFonts repo raw)
    FONT_URLS = {
        "CJK":        "https://github.com/googlefonts/noto-cjk/raw/main/Sans/OTF/SimplifiedChinese/NotoSansSC-Regular.otf",
        "HANGUL":     "https://github.com/googlefonts/noto-cjk/raw/main/Sans/OTF/Korean/NotoSansKR-Regular.otf",
        "ARABIC":     "https://github.com/googlefonts/noto-fonts/raw/main/hinted/ttf/NotoNaskhArabic/NotoNaskhArabic-Regular.ttf",
        "GREEK":      "https://github.com/googlefonts/noto-fonts/raw/main/hinted/ttf/NotoSans/NotoSans-Regular.ttf",
        "OTHER":      "https://github.com/googlefonts/noto-fonts/raw/main/hinted/ttf/NotoSans/NotoSans-Regular.ttf",
        "LATIN": "https://github.com/googlefonts/noto-fonts/raw/main/hinted/ttf/NotoSans/NotoSans-Regular.ttf",
        "DEVANAGARI": "https://github.com/googlefonts/noto-fonts/raw/main/hinted/ttf/NotoSansDevanagari/NotoSansDevanagari-Regular.ttf",
        # CJK heavy: recommend apt-get fonts-noto-cjk instead of downloading single files
    }

    downloaded = []
    for s in scripts:
        if s == "CJK":
            # prefer apt package
            print("CJK script detected. For best results install Noto CJK fonts on your system.")
            print("On Debian/Ubuntu you can run:")
            print("  sudo apt-get update && sudo apt-get install -y fonts-noto-cjk")
            continue
        url = FONT_URLS.get(s, FONT_URLS.get("LATIN"))
        if not url:
            continue
        try:
            fname = os.path.basename(url)
            out_path = os.path.join(fonts_dir, fname)
            if os.path.exists(out_path):
                print(f"Font already present: {out_path}")
                downloaded.append(out_path)
                continue
            import urllib.request
            print("Downloading font:", url)
            urllib.request.urlretrieve(url, out_path)
            os.chmod(out_path, stat.S_IRUSR | stat.S_IWUSR | stat.S_IRGRP | stat.S_IROTH)
            downloaded.append(out_path)
            print(" → saved to", out_path)
        except Exception as e:
            print("Failed to download font:", e)

    # refresh font cache (Linux)
    try:
        print("Refreshing font cache...")
        run(["fc-cache", "-f", "-v"])
    except Exception as e:
        print("fc-cache failed (non-fatal):", e)

    if downloaded:
        print("Fonts installed (user-level). Restart your PDF viewer if needed.")

# -------------------------
# Convert .docx -> .pdf using LibreOffice (preferred)
# -------------------------
def convert_docx_to_pdf_libreoffice(docx_path: str, output_pdf: str) -> str:
    soffice = _which("soffice")
    if not soffice:
        raise RuntimeError(
            "LibreOffice (soffice) not found. Install it for high-fidelity DOCX->PDF.\n"
            "Linux/Colab:  sudo apt-get update && sudo apt-get install -y libreoffice"
        )
    outdir = os.path.dirname(os.path.abspath(output_pdf)) or "."
    _ensure_dir(output_pdf)
    cmd = [soffice, "--headless", "--convert-to", "pdf", "--outdir", outdir, docx_path]
    subprocess.check_call(cmd)
    produced = os.path.join(outdir, os.path.splitext(os.path.basename(docx_path))[0] + ".pdf")
    if not os.path.exists(produced):
        raise RuntimeError("LibreOffice did not produce PDF.")
    if os.path.abspath(produced) != os.path.abspath(output_pdf):
        os.replace(produced, output_pdf)
    return output_pdf

# ---------- Rasterize (guaranteed readability) ----------
def rasterize_pdf_to_pdf(input_pdf: str, output_pdf: str, dpi: int = 200) -> str:
    """
    Converts every page of PDF to an image, then merges them back into PDF.
    Ensures fonts, multilingual text, tables, JSON, etc. are preserved visually.
    """
    _ensure_dir(output_pdf)
    images = convert_from_path(input_pdf, dpi=dpi)
    if not images:
        raise RuntimeError("No pages produced while rasterizing.")

    pdf_pages = []
    for img in images:
        if img.mode != "RGB":
            img = img.convert("RGB")
        pdf_pages.append(img)

    first, rest = pdf_pages[0], pdf_pages[1:]
    first.save(output_pdf, save_all=True, append_images=rest)
    return output_pdf

#==========================
# OCR SETTINGS
#==========================

_EASYOCR_READER = None
_EASYOCR_LANGS = []

def _get_easyocr(langs=['en']):
    global _EASYOCR_READER, _EASYOCR_LANGS
    # Re-initialize if reader is not set or if requested languages are different from currently loaded languages
    if _EASYOCR_READER is None or set(langs) != set(_EASYOCR_LANGS):
        try:
            import easyocr
            # The Reader constructor takes a list of languages directly
            _EASYOCR_READER = easyocr.Reader(langs, gpu=False) # Set gpu=True if you have a GPU
            _EASYOCR_LANGS = langs # Update the stored languages
            print(f"✅ Initialized EasyOCR with languages: {langs}")
        except Exception as e:
            print(f"Error initializing EasyOCR with languages {langs}: {e}")
            _EASYOCR_READER = None
            _EASYOCR_LANGS = [] # Clear languages on failure
    return _EASYOCR_READER

# Preprocessing function (optional but can help)
def preprocess_image(image_path, strong=False):
    img = Image.open(image_path).convert("RGB")
    if strong:
        img = img.filter(ImageFilter.MedianFilter(3))
        enhancer = ImageEnhance.Contrast(img)
        img = enhancer.enhance(2)
        enhancer = ImageEnhance.Brightness(img)
        img = enhancer.enhance(1.5)
    return img

def detect_language_from_image(img):
    """
    Quick heuristic: run a tiny OCR (English + few langs),
    then detect script with langdetect/langid.
    """
    try:
        # Use a basic config for quick detection
        tmp_text = pytesseract.image_to_string(img, lang="eng+hin", config="--psm 6 --oem 3")
        if not tmp_text.strip():
            return "en" # Default to English if no text detected

        # Use langid for better language detection on short text
        code, _ = langid.classify(tmp_text)
        return code
    except Exception as e:
        print(f"Language detection failed: {e}")
        return "en" # Fallback to English on error

def _langs_for_tesseract(easyocr_langs):
    """Maps EasyOCR language codes to Tesseract language codes."""
    tess_map = {
        'en': 'eng', 'hi': 'hin', 'es': 'spa', 'fr': 'fra', 'de': 'deu',
        'ru': 'rus', 'ja': 'jpn', 'ko': 'kor', 'ch_sim': 'chi_sim', 'ch_tra': 'chi_tra'
    }
    tess_langs = [tess_map.get(lang, lang) for lang in easyocr_langs]
    return "+".join(tess_langs)

def _parse_langs_for_easyocr(lang_input):
    """Parses comma-separated or auto language input for EasyOCR."""
    if lang_input.lower() == 'auto':
        # For auto-detection, still need to give EasyOCR *some* languages to load.
        # Let's default to English and Hindi for a common use case, but this could be expanded.
        return ['en', 'hi']
    return [l.strip() for l in lang_input.split(',')]

# =========================
# CSV Converters
# =========================

def csv_to_xls(csv_path, xls_path, chunksize=10000):
    try:
        ensure_parent_dir(xls_path)
        wb = Workbook(write_only=True)
        ws = wb.create_sheet("Sheet1")
        first = True
        for chunk in pd.read_csv(csv_path, chunksize=chunksize, dtype=str):
            if first:
                ws.append(list(chunk.columns))
                first = False
            for row in chunk.itertuples(index=False, name=None):
                ws.append(list(row))
        wb.save(xls_path)
        print(f"✅ CSV → XLSX: {xls_path}")
        return xls_path
    except Exception as e:
        print(f"❌ CSV to XLSX failed: {e}")
        return None

def csv_to_pdf(csv_path, pdf_path, margin=40, line_gap=14, font="Helvetica", size=10):
    try:
        ensure_parent_dir(pdf_path)
        c = canvas.Canvas(pdf_path, pagesize=A4)
        c.setFont(font, size)
        y = PAGE_H - margin
        max_chars = 180
        with open(csv_path, "r", encoding="utf-8", newline="") as f:
            reader = csv.reader(f)
            for row in reader:
                line = " | ".join([str(x) for x in row])
                chunks = [line[i:i+max_chars] for i in range(0, len(line), max_chars)] or [" "]
                for chunk in chunks:
                    if y < margin:
                        c.showPage(); c.setFont(font, size); y = PAGE_H - margin
                    c.drawString(margin, y, chunk)
                    y -= line_gap
        c.save()
        print(f"✅ CSV → PDF: {pdf_path}")
        return pdf_path
    except Exception as e:
        print(f"❌ CSV to PDF failed: {e}")
        return None

def csv_to_doc(csv_path, docx_path):
    try:
        ensure_parent_dir(docx_path)
        df = pd.read_csv(csv_path, dtype=str)
        doc = Document()
        table = doc.add_table(rows=1, cols=len(df.columns))
        for i, col in enumerate(df.columns):
            table.cell(0, i).text = col
        for row in df.itertuples(index=False, name=None):
            cells = table.add_row().cells
            for i, val in enumerate(row):
                cells[i].text = "" if pd.isna(val) else str(val)
        doc.save(docx_path)
        print(f"✅ CSV → DOCX: {docx_path}")
        return docx_path
    except Exception as e:
        print(f"❌ CSV to DOCX failed: {e}")
        return None

def csv_to_txt(csv_path, txt_path, delimiter="\t", chunksize=10000):
    try:
        ensure_parent_dir(txt_path)
        first = True
        with open(txt_path, "w", encoding="utf-8") as out:
            for chunk in pd.read_csv(csv_path, dtype=str, chunksize=chunksize):
                if first:
                    out.write(delimiter.join(chunk.columns) + "\n")
                    first = False
                for row in chunk.itertuples(index=False, name=None):
                    out.write(delimiter.join("" if pd.isna(x) else str(x) for x in row) + "\n")
        print(f"✅ CSV → TXT: {txt_path}")
        return txt_path
    except Exception as e:
        print(f"❌ CSV to TXT failed: {e}")
        return None

def csv_to_json(csv_path, json_path, chunksize=50000):
    try:
        ensure_parent_dir(json_path)
        first = True
        with open(json_path, "w", encoding="utf-8") as out:
            out.write("[")
            for chunk in pd.read_csv(csv_path, dtype=str, chunksize=chunksize):
                records = chunk.where(pd.notnull(chunk), None).to_dict(orient="records")
                for rec in records:
                    if not first: out.write(",\n")
                    json.dump(rec, out, ensure_ascii=False)
                    first = False
            out.write("]")
        print(f"✅ CSV → JSON: {json_path}")
        return json_path
    except Exception as e:
        print(f"❌ CSV to JSON failed: {e}")
        return None

# ===== CSV → Image (single or multi-image ZIP) with correct wrapping =====

def _render_chunk_to_image(lines_chunk, img_path, font_size=12, margin=40, line_height=18, img_width=1200, max_safe_height=30000):
    try:
        # Pre-measure wrapped line count
        total_wrapped_lines = 0
        wrapped_chunks = []
        for line in lines_chunk:
            script = detect_script_simple(line)
            font = get_font(script, font_size)
            wrapped = wrap_text(line, font_size, img_width, margin)
            if not wrapped:
                wrapped = [""]
            wrapped_chunks.append((wrapped, font))
            total_wrapped_lines += len(wrapped)

        img_height = margin * 2 + total_wrapped_lines * line_height
        if img_height > max_safe_height:
            print(f"⚠️ Image height {img_height}px is very large. Consider multi-image ZIP for better reliability.")

        img = Image.new("RGB", (img_width, img_height), "white")
        draw = ImageDraw.Draw(img)

        y = margin
        for wrapped, font in wrapped_chunks:
            for w_line in wrapped:
                draw.text((margin, y), w_line, fill="black", font=font)
                y += line_height

        img.save(img_path, "PNG")
        return img_path
    except Exception as e:
        print(f"❌ Render chunk failed: {e}")
        return None

def csv_to_image(csv_path, out_path, font_size=12, margin=40, line_height=18, img_width=1200, max_lines_per_img=60):
    """
    If CSV has > max_lines_per_img → ask:
      1. Single image
      2. Multi images (ZIP)
    """
    try:
        ensure_parent_dir(out_path)
        lines = []
        with open(csv_path, "r", encoding="utf-8", errors="ignore") as f:
            reader = csv.reader(f)
            for row in reader:
                lines.append(" | ".join(str(x) for x in row))
        if not lines:
            raise RuntimeError("CSV is empty.")

        total = len(lines)
        split = total > max_lines_per_img

        if split:
            print(f"CSV lines: {total} (> {max_lines_per_img})")
            print("Choose output option:")
            print("1. Single giant image (may be very tall)")
            print("2. Multiple images (zipped)")
            choice = input("Enter choice (1/2): ").strip()
            split = (choice != "1")

        images = []
        if not split:
            if not out_path.lower().endswith(".png"):
                out_path += ".png"
            path = _render_chunk_to_image(lines, out_path, font_size, margin, line_height, img_width)
            if path: images.append(path)
            print(f"✅ CSV → Image: {out_path}")
            return out_path
        else:
            chunks = [lines[i:i+max_lines_per_img] for i in range(0, total, max_lines_per_img)]
            img_dir = os.path.splitext(out_path)[0] + "_images"
            os.makedirs(img_dir, exist_ok=True)
            for i, chunk in enumerate(chunks):
                img_path = os.path.join(img_dir, f"page_{i+1}.png")
                _render_chunk_to_image(chunk, img_path, font_size, margin, line_height, img_width)
                images.append(img_path)
            zip_path = os.path.splitext(out_path)[0] + ".zip"
            with zipfile.ZipFile(zip_path, 'w') as zf:
                for img_file in images:
                    zf.write(img_file, os.path.basename(img_file))
            shutil.rmtree(img_dir)
            print(f"✅ CSV → Images ZIP: {zip_path}")
            return zip_path
    except Exception as e:
        print(f"❌ CSV to Image failed: {e}")
        return None

# =========================
# XLS Converters
# =========================

def xls_to_csv(xls_path, csv_path, sheet_name=None):
    try:
        ensure_parent_dir(csv_path)
        wb = load_workbook(xls_path, read_only=True, data_only=True)
        ws = wb[sheet_name] if sheet_name else wb.active
        with open(csv_path, "w", encoding="utf-8", newline="") as out:
            writer = csv.writer(out)
            for row in ws.iter_rows(values_only=True):
                writer.writerow(["" if v is None else v for v in row])
        print(f"✅ XLSX → CSV: {csv_path}")
        return csv_path
    except Exception as e:
        print(f"❌ XLSX to CSV failed: {e}")
        return None

def xls_to_doc(xls_path, docx_path, sheet_name=None):
    try:
        ensure_parent_dir(docx_path)
        wb = load_workbook(xls_path, read_only=True, data_only=True)
        ws = wb[sheet_name] if sheet_name else wb.active
        doc = Document()
        rows = list(ws.iter_rows(values_only=True))
        if not rows:
            doc.save(docx_path); print(f"⚠️ Sheet empty. DOCX saved: {docx_path}"); return docx_path
        table = doc.add_table(rows=1, cols=len(rows[0]))
        for i, val in enumerate(rows[0]):
            table.cell(0,i).text = "" if val is None else str(val)
        for r in rows[1:]:
            cells = table.add_row().cells
            for i, val in enumerate(r):
                cells[i].text = "" if val is None else str(val)
        doc.save(docx_path)
        print(f"✅ XLSX → DOCX: {docx_path}")
        return docx_path
    except Exception as e:
        print(f"❌ XLSX to DOCX failed: {e}")
        return None

def xls_to_txt(xls_path, txt_path, delimiter="\t", sheet_name=None):
    try:
        ensure_parent_dir(txt_path)
        wb = load_workbook(xls_path, read_only=True, data_only=True)
        ws = wb[sheet_name] if sheet_name else wb.active
        with open(txt_path, "w", encoding="utf-8") as out:
            for row in ws.iter_rows(values_only=True):
                out.write(delimiter.join("" if v is None else str(v) for v in row) + "\n")
        print(f"✅ XLSX → TXT: {txt_path}")
        return txt_path
    except Exception as e:
        print(f"❌ XLSX to TXT failed: {e}")
        return None

def xls_to_pdf(xls_path, pdf_path, margin=40, line_gap=14, font="Helvetica", size=10, sheet_name=None):
    try:
        ensure_parent_dir(pdf_path)
        wb = load_workbook(xls_path, read_only=True, data_only=True)
        ws = wb[sheet_name] if sheet_name else wb.active
        c = canvas.Canvas(pdf_path, pagesize=A4)
        c.setFont(font, size)
        y = PAGE_H - margin
        max_chars = 180
        for row in ws.iter_rows(values_only=True):
            line = " | ".join("" if v is None else str(v) for v in row)
            chunks = [line[i:i+max_chars] for i in range(0, len(line), max_chars)] or [" "]
            for ch in chunks:
                if y < margin:
                    c.showPage(); c.setFont(font, size); y = PAGE_H - margin
                c.drawString(margin, y, ch)
                y -= line_gap
        c.save()
        print(f"✅ XLSX → PDF: {pdf_path}")
        return pdf_path
    except Exception as e:
        print(f"❌ XLSX to PDF failed: {e}")
        return None

def xls_to_json(xls_path, json_path, sheet_name=None):
    try:
        ensure_parent_dir(json_path)
        wb = load_workbook(xls_path, read_only=True, data_only=True)
        ws = wb[sheet_name] if sheet_name else wb.active
        rows = list(ws.iter_rows(values_only=True))
        if not rows:
            with open(json_path, "w", encoding="utf-8") as out:
                out.write("[]")
            print(f"⚠️ Sheet empty. JSON saved: {json_path}")
            return json_path
        headers = [str(h) if h is not None else "" for h in rows[0]]
        data = []
        for r in rows[1:]:
            obj = {headers[i]: ("" if r[i] is None else r[i]) for i in range(len(headers))}
            data.append(obj)
        with open(json_path, "w", encoding="utf-8") as out:
            json.dump(data, out, ensure_ascii=False, indent=2)
        print(f"✅ XLSX → JSON: {json_path}")
        return json_path
    except Exception as e:
        print(f"❌ XLSX to JSON failed: {e}")
        return None


def xls_to_image(xls_path, out_path, sheet_name=None, font_size=12, margin=40, line_height=18,
                 img_width=1200, max_lines_per_img=60, max_safe_height=30000):
    """
    Converts an XLS/XLSX file to PNG image(s).
    Handles large files by splitting into multiple images.
    """
    try:
        if not os.path.exists(xls_path):
            raise FileNotFoundError(f"❌ File not found: {xls_path}")

        wb = load_workbook(xls_path, read_only=True, data_only=True)
        ws = wb[sheet_name] if sheet_name else wb.active

        lines = []
        for row in ws.iter_rows(values_only=True):
            lines.append(" | ".join(str(x) if x is not None else "" for x in row))

        if not lines:
            raise RuntimeError("❌ XLS/XLSX sheet is empty.")

        total_lines = len(lines)
        split_into_multiple = total_lines > max_lines_per_img

        # Ask user if too many lines
        if split_into_multiple:
            print(f"Sheet has {total_lines} lines (>{max_lines_per_img}).")
            print("Choose output option:")
            print("1. Single giant image (may be very tall)")
            print("2. Multiple images (zipped)")
            choice = input("Enter choice (1/2): ").strip()
            if choice == "1":
                split_into_multiple = False
            else:
                split_into_multiple = True

        images = []

        def render_chunk(chunk, filename):
            """Render one chunk of XLS data into an image file."""
            # --- Measure wrapped line count first ---
            total_wrapped_lines = 0
            wrapped_chunks = []
            for line in chunk:
                script = detect_script_simple(line)
                font = get_font_path(script, font_size)
                wrapped = textwrap.wrap(line, width=int((img_width - 2*margin) / (font_size*0.6)))
                if not wrapped:
                    wrapped = [""]
                wrapped_chunks.append((wrapped, font))
                total_wrapped_lines += len(wrapped)

            # --- Now allocate correct height ---
            img_height = margin * 2 + total_wrapped_lines * line_height
            if img_height > max_safe_height:
                print(f"⚠️ Warning: Image height {img_height}px may be too large.")
            img = Image.new("RGB", (img_width, img_height), "white")
            draw = ImageDraw.Draw(img)

            y_offset = margin
            for wrapped, font in wrapped_chunks:
                for w_line in wrapped:
                    draw.text((margin, y_offset), w_line, fill="black", font=font)
                    y_offset += line_height

            img.save(filename, "PNG")
            return filename


        if not split_into_multiple:
            # --- Option 1: single image ---
            if not out_path.lower().endswith(".png"):
                out_path += ".png"
            images.append(render_chunk(lines, out_path))

        else:
            # --- Option 2: multiple images zipped ---
            chunks = [lines[i:i+max_lines_per_img] for i in range(0, total_lines, max_lines_per_img)]
            img_dir = os.path.splitext(out_path)[0] + "_images"
            os.makedirs(img_dir, exist_ok=True)

            for i, chunk in enumerate(chunks):
                img_path = os.path.join(img_dir, f"page_{i+1}.png")
                images.append(render_chunk(chunk, img_path))

            # Zip all pages
            zip_path = os.path.splitext(out_path)[0] + ".zip"
            with zipfile.ZipFile(zip_path, 'w') as zf:
                for img_file in images:
                    zf.write(img_file, os.path.basename(img_file))

            shutil.rmtree(img_dir)  # cleanup
            images = [zip_path]
            out_path = zip_path


        print(f"✅ XLS/XLSX converted to image(s): {out_path}")
        return out_path

    except Exception as e:
        print(f"❌ XLS/XLSX to Image conversion failed: {e}")
        return None


# =========================
# TXT Converters
# =========================

def txt_to_pdf(txt_path, pdf_path, font_size=12, margin=40, line_gap=16):

    # Register multiple fonts
    pdfmetrics.registerFont(TTFont("NotoSans", "/content/NotoSans-Regular.ttf")) #Add Local path of font
    pdfmetrics.registerFont(TTFont("NotoSansDevanagari", "/content/NotoSansDevanagari-Medium.ttf")) #Add Local path of font

    try:
        c = canvas.Canvas(pdf_path, pagesize=A4)

        y = PAGE_H - margin
        max_width = PAGE_W - 2 * margin

        with open(txt_path, "r", encoding="utf-8", errors="ignore") as f:
            for raw_line in f:
                line = raw_line.strip("\n")

                # choose font dynamically
                font_name = detect_script_text(line)
                c.setFont(font_name, font_size)

                wrapped_lines = simpleSplit(line, font_name, font_size, max_width)
                if not wrapped_lines:
                    wrapped_lines = [" "]

                for wline in wrapped_lines:
                    if y < margin:  # new page
                        c.showPage()
                        y = PAGE_H - margin
                        c.setFont(font_name, font_size)

                    c.drawString(margin, y, wline)
                    y -= line_gap

        c.save()
        print(f"✅ PDF saved: {pdf_path}")
        return pdf_path
    except Exception as e:
        print("❌ txt_to_pdf failed:", str(e))
        return None

def txt_to_doc(input_file, output_file):
    try:
        if not txt_path.endswith(".txt"):
            raise RuntimeError("❌ Only TEXT files are supported!")

        if not os.path.exists(txt_path):
            raise RuntimeError("❌ File not found!")

        with open(input_file, "r", encoding="utf-8", errors="ignore") as f:
            text = f.read()

        # ✅ Detect language
        try:
            language = detect(text)
        except:
            language = "en"

        # ✅ Font mapping by language
        font_map = {
            #"hi": "Mangal",        # Hindi
            "hi": "Nirmala UI",     # Hindi
            "bn": "Vrinda",           # Bengali
            "ta": "Latha",            # Tamil
            "te": "Gautami",          # Telugu
            "kn": "Tunga",            # Kannada
            "ml": "Kartika",          # Malayalam
            "gu": "Shruti",           # Gujarati
            "pa": "Raavi",            # Punjabi
            "or": "Kalinga",          # Odia

            "zh-cn": "Microsoft YaHei",   # Simplified Chinese
            "zh-tw": "PMingLiU",          # Traditional Chinese
            "ja": "MS Mincho",            # Japanese
            "ko": "Malgun Gothic",        # Korean

            "ar": "Traditional Arabic",   # Arabic
            "fa": "B Nazanin",            # Persian/Farsi
            "ur": "Jameel Noori Nastaleeq", # Urdu
            "he": "David",                # Hebrew

            "en": "Arial",         # English
            "ru": "Times New Roman",      # Russian
            "uk": "Times New Roman",      # Ukrainian
            "el": "Times New Roman",      # Greek
            "th": "Angsana New",          # Thai
            "vi": "Times New Roman",      # Vietnamese

            "fr": "Calibri",              # French
            "de": "Calibri",              # German
            "es": "Calibri",              # Spanish
            "it": "Calibri",              # Italian
            "pt": "Calibri",              # Portuguese
        }
        chosen_font = font_map.get(language, "Noto Sans") # fallback Noto Sans

        doc = Document()
        style = doc.styles['Normal']
        style.font.name = chosen_font
        style.font.size = Pt(12)

        for line in text.splitlines():
            if line.strip():  # skip blank lines
                    p = doc.add_paragraph(line.strip())
                    r = p.runs[0]
                    r.font.name = chosen_font
                    r._element.rPr.rFonts.set(qn('w:eastAsia'), chosen_font)

        doc.save(output_file)
        print(f"✅ Saved {output_file} with font '{chosen_font}' (detected lang: {language})")

        # ✅ Suggestion for user
        if language not in font_map:
            print(f"⚠️ Language '{language}' not mapped. "
                  f"Please try 'Noto Sans' or manually set an appropriate font in MS Word.")
        else:
            print(f"💡 Tip: If the text doesn’t render well, "
                  f"manually set the font in MS Word to '{chosen_font}' for proper display.")

    except Exception as e:
        print(f"❌ TXT to DOC failed: {e}")
        return None

def txt_to_json(txt_path, json_path):
    try:
        ensure_parent_dir(json_path)
        if not txt_path.endswith(".txt"):
            raise RuntimeError("❌ Only TEXT files are supported!")

        if not os.path.exists(txt_path):
            raise RuntimeError("❌ File not found!")

        data = []
        with open(txt_path, "r", encoding="utf-8", errors="ignore") as f:
            for line in f:
                s = line.strip()
                if s:
                    data.append(s)
        with open(json_path, "w", encoding="utf-8") as out:
            json.dump(data, out, ensure_ascii=False, indent=2)
        return json_path

    except Exception as e:
        print(f"❌ TXT to JSON failed: {e}")
        return None

def txt_to_csv(txt_path, csv_path, delimiter="\t", auto_detect_delimiter=False,
    excel_friendly=True,   # writes utf-8-sig so Excel auto-detects BOM
    notify_user=True,
    save_xlsx=False        # if True and pandas available, also write .xlsx beside csv
):
    """
    Stream text -> CSV preserving UTF-8 and user friendly Excel behavior.

    Parameters:
      - txt_path: input text file path (UTF-8 expected)
      - csv_path: output path (.csv)
      - delimiter: delimiter to split lines into columns (default tab)
      - auto_detect_delimiter: try to detect delimiter automatically from file sample
      - excel_friendly: if True write using encoding='utf-8-sig' (BOM) -> Excel on Windows reads well
      - notify_user: print or popup instructions on how to open in Excel if non-English content detected
      - save_xlsx: if True and pandas installed, also write an .xlsx copy (recommended)
    """
    try:
        if not os.path.exists(txt_path):
            raise FileNotFoundError(f"Input file not found: {txt_path}")

        # auto-detect delimiter if requested
        if auto_detect_delimiter:
            try:
                delimiter = _auto_detect_delimiter(txt_path)
            except Exception:
                pass

        # detect non-english content & scripts
        non_english, top_langs, has_non_ascii, scripts = _detect_non_english_and_scripts(txt_path)

        # choose encoding
        enc = "utf-8-sig" if excel_friendly else "utf-8"

        # Write CSV streaming line-by-line
        with open(txt_path, "r", encoding="utf-8", errors="ignore") as src, \
             open(csv_path, "w", encoding=enc, newline="") as dst:
            writer = csv.writer(dst)
            for raw in src:
                row = raw.rstrip("\n").split(delimiter)
                # ensure row is list of str (no bytes)
                row = [("" if v is None else str(v)) for v in row]
                writer.writerow(row)

        # Optionally save XLSX to avoid Excel encoding issues completely
        xlsx_path = None
        if save_xlsx:
            if not _HAS_PANDAS:
                print("⚠️ pandas not installed; cannot save .xlsx. Install pandas + openpyxl to enable this.")
            else:
                try:
                    # read via pandas (fast for reasonably sized files)
                    df = pd.read_csv(csv_path, encoding=enc)
                    xlsx_path = os.path.splitext(csv_path)[0] + ".xlsx"
                    df.to_excel(xlsx_path, index=False, engine="openpyxl")
                except Exception as e:
                    print("⚠️ Could not create .xlsx:", e)

        # Notify user (print or popup) with instructions if non-english content or user asked to be notified
        if notify_user:
            _notify_user_excel_instructions(csv_path, non_english, top_langs, scripts)

        return {"csv": csv_path, "xlsx": xlsx_path, "non_english": non_english, "langs": top_langs, "scripts": scripts}

    except Exception as e:
        print(f"❌ TXT to CSV failed: {e}")
        return None

def txt_to_image(txt_path, output_dir, font_size=24, width=1240, height=1754, margin=40, max_lines_per_img=70, split=None):

    try:
        if not txt_path.endswith(".txt"):
            raise RuntimeError("❌ Only TEXT files are supported!")

        if not os.path.exists(txt_path):
            raise RuntimeError("❌ File not found!")

        # Read text file
        with open(txt_path, "r", encoding="utf-8", errors="ignore") as f:
            lines = f.readlines()

        os.makedirs(output_dir, exist_ok=True)
        images = []
        line_height = font_size + 10

        # Auto decide if split is None
        if split is None:
            if len(lines) <= 20:
                split = False   # single image
            else:
                # user choice
                choice = input("Text has more than 20 lines. Do you want multiple images? (y/n): ")
                split = True if choice.lower().startswith("y") else False

        if not split:
            # Option 1: All lines in ONE big image
            total_height = margin*2 + len(lines)*line_height
            img = Image.new("RGB", (width, total_height), "white")
            draw = ImageDraw.Draw(img)

            y = margin
            for line in lines:
                font = get_font_for_line(line, font_size)
                draw.text((margin, y), line.strip(), fill="black", font=font)
                y += line_height

            out_path = os.path.join(output_dir, "output_single.png")
            img.save(out_path)
            images.append(out_path)

        else:
            # Option 2: Split into multiple images
            chunks = [lines[i:i+max_lines_per_img] for i in range(0, len(lines), max_lines_per_img)]

            for idx, chunk in enumerate(chunks, start=1):
                img = Image.new("RGB", (width, height), "white")
                draw = ImageDraw.Draw(img)

                y = margin
                for line in chunk:
                    font = get_font_for_line(line, font_size)
                    draw.text((margin, y), line.strip(), fill="black", font=font)
                    y += line_height

                out_path = os.path.join(output_dir, f"page_{idx}.png")
                img.save(out_path)
                images.append(out_path)

        print(f"✅ Images saved: {images}")
        return images

    except Exception as e:
        print("❌ TXT to Image failed:", str(e))
        return None

def txt_to_xls(txt_path, xls_path, delimiter="\t"):
    try:
        if not txt_path.endswith(".txt"):
            raise RuntimeError("❌ Only TEXT files are supported!")

        if not os.path.exists(txt_path):
            raise RuntimeError("❌ File not found!")

        wb = Workbook(write_only=True)
        ws = wb.create_sheet("Sheet1")
        with open(txt_path, "r", encoding="utf-8", errors="ignore") as src:
            for line in src:
                ws.append(line.rstrip("\n").split(delimiter))
        wb.save(xls_path)
        return xls_path

    except Exception as e:
        print("❌ TXT to XLS failed:", str(e))
        return None

# =========================
# JSON Converters
# =========================

def json_to_csv(json_path, csv_path):
    """
    Convert JSON or NDJSON into CSV.
    - Handles large files (50MB+).
    - UTF-8 safe (Hindi + English).
    - Supports list-of-objects JSON & NDJSON.
    """
    try:
        if not json_path.endswith(".json"):
            raise RuntimeError("❌ Only JSON files are supported!")

        if not os.path.exists(json_path):
            raise RuntimeError("❌ File not found!")

        with open(csv_path, "w", encoding="utf-8-sig", newline="") as out:
            # utf-8-sig ensures Excel also reads Hindi properly
            writer = None

            try:
                # ---- Try as normal JSON (array of objects) ----
                with open(json_path, "r", encoding="utf-8") as f:
                    data = json.load(f)

                if isinstance(data, dict):
                    # if root is dict → wrap into list
                    data = [data]

                # Ensure it's list of dicts
                if not isinstance(data, list):
                    raise ValueError("JSON is not a list-of-objects")

                # Write CSV
                for idx, obj in enumerate(data):
                    if writer is None:
                        writer = csv.DictWriter(out, fieldnames=list(obj.keys()))
                        writer.writeheader()
                    writer.writerow(obj)

            except json.JSONDecodeError:
                # ---- NDJSON fallback (streaming) ----
                with open(json_path, "r", encoding="utf-8") as f:
                    for obj in ijson.items(f, "", multiple_values=True):
                        if isinstance(obj, dict):
                            if writer is None:
                                writer = csv.DictWriter(out, fieldnames=list(obj.keys()))
                                writer.writeheader()
                            writer.writerow(obj)
                        else:
                            # if NDJSON line is not dict
                            if writer is None:
                                writer = csv.writer(out)
                                writer.writerow(["value"])
                            writer.writerow([obj])

        return csv_path

    except Exception as e:
        print(f"❌ JSON to CSV failed: {e}")
        return None

def json_to_xls(json_path, xls_path):
    """
    Convert JSON or NDJSON into XLSX.
    - Handles large files via streaming
    - UTF-8 safe (Hindi/English)
    - Flattens nested JSON
    """
    try:
        if not json_path.endswith(".json"):
            raise RuntimeError("❌ Only JSON files are supported!")

        if not os.path.exists(json_path):
            raise RuntimeError("❌ File not found!")

        wb = Workbook(write_only=True)
        ws = wb.create_sheet("Sheet1")

        headers_written = False
        headers = None

        with open(json_path, "r", encoding="utf-8") as f:
            try:
                # ---- simple JSON ----
                data = json.load(f)
                if isinstance(data, dict):
                    data = [data]
                if not isinstance(data, list):
                    raise ValueError("JSON is not a list of objects.")

                for obj in data:
                    if not isinstance(obj, dict):
                        continue
                    flat = flatten_json(obj)
                    if not headers_written:
                        headers = list(flat.keys())
                        ws.append(headers)
                        headers_written = True
                    ws.append([flat.get(h, "") for h in headers])

            except json.JSONDecodeError:
                # ---- NDJSON ----
                f.seek(0)
                for obj in ijson.items(f, "", multiple_values=True):
                    if not isinstance(obj, dict):
                        continue
                    flat = flatten_json(obj)
                    if not headers_written:
                        headers = list(flat.keys())
                        ws.append(headers)
                        headers_written = True
                    ws.append([flat.get(h, "") for h in headers])

        wb.save(xls_path)
        wb.close()

        if not os.path.exists(xls_path):
            return None
        return xls_path

    except Exception as e:
        print(f"❌ JSON to XLSX failed: {e}")
        return None

def json_to_pdf(json_path, output_pdf):
    try:

        # --- Only allow .json files ---
        if not json_path.lower().endswith(".json"):
            print(f"❌ Unsupported file type: {json_path}. Only .json files are supported.")

        if not os.path.exists(json_path):
            raise RuntimeError("Input file not found: " + json_path)

        # Detect JSON type: normal vs JSON lines
        data_list = []
        with open(json_path, "r", encoding="utf-8") as f:
            first_char = f.read(1)
            f.seek(0)
            if first_char == "{":
                # Could be JSON Lines or single JSON object
                try:
                    # Try loading whole file as a JSON object
                    data = json.load(f)
                    if isinstance(data, dict):
                        data_list.append(data)
                    elif isinstance(data, list):
                        data_list = data
                except json.JSONDecodeError:
                    # Treat as JSON Lines
                    f.seek(0)
                    for line in f:
                        line = line.strip()
                        if line:
                            data_list.append(json.loads(line))
            elif first_char == "[":
                data_list = json.load(f)
            else:
                raise RuntimeError("Invalid JSON format")

        # Convert to pretty JSON string
        json_text = json.dumps(data_list, indent=4, ensure_ascii=False)

        # HTML content for PDF
        html_content = f"""
        <html>
        <head>
            <meta charset="utf-8">
            <style>
                body {{ font-family: 'Noto Sans', 'Devanagari Sans', monospace; font-size:12pt; }}
                pre {{ white-space: pre-wrap; word-wrap: break-word; }}
            </style>
        </head>
        <body>
            <pre>{json_text}</pre>
        </body>
        </html>
        """

        tmp_html = tempfile.mktemp(suffix=".html")
        with open(tmp_html, "w", encoding="utf-8") as f:
            f.write(html_content)

        pdfkit.from_file(tmp_html, output_pdf)
        os.remove(tmp_html)
        return output_pdf

    except Exception as e:
        print(f"❌ JSON to PDF failed: {e}")
        return None

def json_to_txt(json_path, txt_path):
    """
    Convert JSON/NDJSON to TXT.
    - Handles large files (50MB+).
    - Simple JSON => pretty indented text.
    - NDJSON => record-by-record streaming.
    - UTF-8 safe (Hindi + English).
    """
    try :

        if not json_path.endswith(".json"):
              raise RuntimeError("❌ Only JSON files are supported!")

        if not os.path.exists(json_path):
              raise RuntimeError("❌ File not found!")

        with open(json_path, "r", encoding="utf-8") as f, \
            open(txt_path, "w", encoding="utf-8") as out:

            try:
                # Try loading as a whole JSON (simple JSON)
                data = json.load(f)
                # pretty print
                pretty = json.dumps(data, ensure_ascii=False, indent=4)
                out.write(pretty + "\n")

            except json.JSONDecodeError:
                # Fallback: NDJSON mode
                f.seek(0)
                line_no = 0
                for line in f:
                    line_no += 1
                    line = line.strip()
                    if not line:
                        continue
                    try:
                        obj = json.loads(line)
                        # pretty-print each object with indentation
                        pretty = json.dumps(obj, ensure_ascii=False, indent=4)
                        out.write(f"--- Record {line_no} ---\n{pretty}\n\n")
                    except Exception:
                        out.write(f"⚠️ Skipping bad line {line_no}: {line[:50]}\n")

        return txt_path

    except Exception as e:
        print(f"❌ JSON to TXT failed: {e}")
        return None

def json_to_doc(json_file, output_file):
    try:
        if not json_file.endswith(".json"):
            raise RuntimeError("❌ Only JSON files are supported!")

        if not os.path.exists(json_file):
            raise RuntimeError("❌ File not found!")

        data_list = []
        with open(json_file, "r", encoding="utf-8") as f:
            first_char = f.read(1)
            f.seek(0)
            if first_char == "{":
                try:
                    # Normal JSON
                    data = json.load(f)
                    if isinstance(data, dict):
                        data_list.append(data)
                    elif isinstance(data, list):
                        data_list = data
                except json.JSONDecodeError:
                    # NDJSON
                    f.seek(0)
                    for line in f:
                        line = line.strip()
                        if line:
                            data_list.append(json.loads(line))
            elif first_char == "[":
                data_list = json.load(f)
            else:
                raise RuntimeError("❌ Invalid JSON format")

        if not data_list:
            raise RuntimeError("❌ No data found in JSON.")

        # --- Create DOCX ---
        doc = Document()
        style = doc.styles['Normal']
        font = style.font
        font.name = 'Nirmala UI'   # Hindi-friendly font (fallback)
        font.size = Pt(11)
        style._element.rPr.rFonts.set(qn('w:eastAsia'), 'Nirmala UI')

        doc.add_heading("JSON Data Export", level=1)

        # --- Output formatting ---
        if isinstance(data_list, list) and all(isinstance(item, dict) for item in data_list):
            # Multiple dicts → print each as formatted block
            for idx, obj in enumerate(data_list, 1):
                doc.add_paragraph(f"Record {idx}:", style='Heading 2')
                pretty = json.dumps(obj, indent=4, ensure_ascii=False)
                doc.add_paragraph(pretty)
        else:
            # Single object or nested → dump as pretty JSON
            pretty = json.dumps(data_list, indent=4, ensure_ascii=False)
            doc.add_paragraph(pretty)

        # --- Save DOCX ---
        doc.save(output_file)
        return output_file

    except Exception as e:
        print(f"❌ JSON to DOC failed: {e}")
        return None

def json_to_image(json_path, output_path):
    try:
        if not json_path.endswith(".json"):
            raise RuntimeError("❌ Only JSON files are supported!")

        if not os.path.exists(json_path):
            raise RuntimeError("❌ File not found!")

        ndjson_mode = is_ndjson(json_path)

        print("Choose conversion mode:")
        print("1. Single big image (whole JSON in one PNG)")
        print("2. Split into 30-line chunks (ZIP of PNGs)")
        choice = input("Enter choice (1/2): ").strip()

        if ndjson_mode:
            print("📂 NDJSON detected → routing via json_to_pdf()...")
            tmp_pdf = tempfile.mktemp(suffix=".pdf")
            json_to_pdf(json_path, tmp_pdf)   # use your fast code

            # convert PDF → images
            images = convert_from_path(tmp_pdf, dpi=100)

            if choice == "1":
                # merge pages into one long PNG
                total_height = sum(i.height for i in images)
                max_width = max(i.width for i in images)
                big_img = Image.new("RGB", (max_width, total_height), "white")
                y_offset = 0
                for img in images:
                    big_img.paste(img, (0, y_offset))
                    y_offset += img.height
                big_img.save(output_path, "PNG")
                print(f"✅ NDJSON → single long image: {output_path}")
                return output_path
            else:
                # each page separate → zip
                tmpdir = tempfile.mkdtemp()
                img_files = []
                base, ext = os.path.splitext(output_path)
                for i, page in enumerate(images, 1):
                    out_file = os.path.join(tmpdir, f"page_{i}.png")
                    page.save(out_file, "PNG")
                    img_files.append(out_file)

                zip_path = output_path.replace(".png", ".zip")
                with zipfile.ZipFile(zip_path, "w") as zipf:
                    for img in img_files:
                        zipf.write(img, os.path.basename(img))

                print(f"✅ NDJSON → images zipped: {zip_path}")
                return zip_path

        else:
            # ✅ Normal JSON → existing HTML flow
            with open(json_path, "r", encoding="utf-8") as f:
                data_json = json.load(f)
            if isinstance(data_json, dict):
                data_json = [data_json]
            pretty_json = json.dumps(data_json, indent=4, ensure_ascii=False)

            if choice == "1":
                return save_html_as_image(pretty_json, output_path, long_mode=True)
            else:
                lines = pretty_json.split("\n")
                tmpdir = tempfile.mkdtemp()
                img_files = []
                chunk_size = 30
                for i in range(0, len(lines), chunk_size):
                    chunk = "\n".join(lines[i:i + chunk_size])
                    img_file = os.path.join(tmpdir, f"chunk_{i//chunk_size + 1}.png")
                    save_html_as_image(chunk, img_file, long_mode=False)
                    img_files.append(img_file)

                zip_path = output_path.replace(".png", ".zip")
                with zipfile.ZipFile(zip_path, "w") as zipf:
                    for img in img_files:
                        zipf.write(img, os.path.basename(img))

                print(f"✅ Converted successfully (chunks zipped): {zip_path}")
                return zip_path

    except Exception as e:
        print(f"❌ JSON to XLSX failed: {e}")
        return None

def save_html_as_image(text, out_path, long_mode=False):
    html_content = f"""
    <html>
    <head>
        <meta charset="UTF-8">
        <style>
            body {{
                font-family: 'Nirmala UI','Noto Sans Devanagari','Noto Sans','Arial Unicode MS',sans-serif;
                font-size: 14pt;
                white-space: pre-wrap;
            }}
            pre {{
                white-space: pre-wrap;
                word-wrap: break-word;
            }}
        </style>
    </head>
    <body><pre>{text}</pre></body></html>
    """
    tmp_html = tempfile.mktemp(suffix=".html")
    with open(tmp_html, "w", encoding="utf-8") as f:
        f.write(html_content)

    tmp_pdf = out_path.replace(".png", f"_{uuid.uuid4().hex}.pdf")
    HTML(tmp_html).write_pdf(tmp_pdf)
    images = convert_from_path(tmp_pdf, dpi=100)

    if long_mode and len(images) > 1:
        total_height = sum(i.height for i in images)
        max_width = max(i.width for i in images)
        big_img = Image.new("RGB", (max_width, total_height), "white")
        y_offset = 0
        for img in images:
            big_img.paste(img, (0, y_offset))
            y_offset += img.height
        big_img.save(out_path, "PNG")
        print(f"✅ Saved single long image: {out_path}")
        return out_path
    else:
        images[0].save(out_path, "PNG")
        print(f"✅ Saved image: {out_path}")
        return out_path

# =========================
# PDF Helpers
# =========================

def pdf_to_txt(pdf_path, txt_path, ocr_langs="eng+hin+jpn+chi_sim+chi_tra+deu+fra"):
    '''Support English + Hindi + Japanese + Chinese (simplified + traditional)
    *Note : sudo apt-get install tesseract-ocr \
     tesseract-ocr-hin \
     tesseract-ocr-jpn \
     tesseract-ocr-chi-sim \
     tesseract-ocr-chi-tra \
     tesseract-ocr-deu \
     tesseract-ocr-fra
     install languange pack for better experience 
    '''
    try:
        ensure_parent_dir(txt_path)
        if not pdf_path.endswith(".pdf"):
            raise RuntimeError("❌ Only PDF files are supported!")
        if not os.path.exists(pdf_path):
            raise RuntimeError("❌ File not found!")

        extracted_text = ""

        # Try direct text extraction
        with pdfplumber.open(pdf_path) as pdf:
            for page in pdf.pages:
                text = page.extract_text()
                if text and text.strip():
                    extracted_text += text.rstrip() + "\n\n"

        # Always run OCR as supplement
        images = convert_from_path(pdf_path)
        ocr_text = ""
        for img in images:
            ocr_text += pytesseract.image_to_string(img, lang=ocr_langs) + "\n"

        # Combine both
        final_text = (extracted_text + "\n" + ocr_text).strip()

        with open(txt_path, "w", encoding="utf-8") as out:
            out.write(final_text)

        return txt_path
    except Exception as e:
        print(f"❌ PDF to TXT failed: {e}")
        return None

def pdf_to_txt_ocr(pdf_path, txt_path, lang="eng+hin"):
    '''Support English + Hindi'''
    try:
        images = convert_from_path(pdf_path)
        full_text = ""
        for img in images:
            text = pytesseract.image_to_string(img, lang=lang)
            full_text += text + "\n"

        # Ensure the output directory exists
        output_dir = os.path.dirname(txt_path)
        if output_dir and not os.path.exists(output_dir):
            os.makedirs(output_dir)

        with open(txt_path, "w", encoding="utf-8") as f:
            f.write(full_text)
        return txt_path
    except Exception as e:
        print(f"❌ OCR failed: {e}")
        return None

#def pdf_to_txt(pdf_path, txt_path):
#    try:
#        ensure_parent_dir(txt_path)
#        text = pdf_extract_text(pdf_path) or ""
#        with open(txt_path, "w", encoding="utf-8") as out:
#            out.write(text)
#        print(f"✅ PDF → TXT: {txt_path}")
#        return txt_path
#    except Exception as e:
#        print(f"❌ PDF to TXT failed: {e}")
#        return None

def pdf_to_docx(pdf_path, docx_path):
    try:
        if not pdf_path.endswith(".pdf"):
            raise RuntimeError("❌ Only PDF files are supported!")

        if not os.path.exists(pdf_path):
            raise RuntimeError("❌ File not found!")

        doc = Document()

        with pdfplumber.open(pdf_path) as pdf:
            for pageno, page in enumerate(pdf.pages, 1):

                # --- Extract Text ---
                text = page.extract_text() or ""
                for line in text.splitlines():
                    script = detect_script_pdf(line)
                    font_name = FALLBACK_FONTS_PDF.get(script, FALLBACK_FONTS_PDF["DEFAULT"])
                    para = doc.add_paragraph()
                    run = para.add_run(line)
                    run.font.name = font_name
                    run.font.size = Pt(11)

                    # set for East Asian / Complex Scripts
                    rPr = run._element.rPr.rFonts
                    rPr.set(qn('w:eastAsia'), font_name)
                    rPr.set(qn('w:cs'), font_name)
                    rPr.set(qn('w:hAnsi'), font_name)

                # --- Extract Images ---
                for img in page.images:
                    try:
                        x0, top, x1, bottom = img["x0"], img["top"], img["x1"], img["bottom"]
                        cropped = page.crop((x0, top, x1, bottom))
                        pil_img = cropped.to_image(resolution=150).original
                        img_path = f"temp_img_{pageno}.png"
                        pil_img.save(img_path)

                        # insert into docx
                        doc.add_picture(img_path, width=Inches(4))
                    except Exception as e:
                        print(f"⚠️ Image on page {pageno} skipped: {e}")

                if pageno < len(pdf.pages):
                    doc.add_page_break()

        doc.save(docx_path)
        return docx_path

    except Exception as e:
        print(f"❌ PDF to DOC failed: {e}")
        return None

def pdf_to_image(pdf_path, out_dir, fmt="png", dpi=150, base_name="page"):
    try:
        os.makedirs(out_dir, exist_ok=True)
        if not pdf_path.endswith(".pdf"):
            raise RuntimeError("❌ Only PDF files are supported!")

        if not os.path.exists(pdf_path):
            raise RuntimeError("❌ File not found!")

        # convert all pages
        images = convert_from_path(pdf_path, dpi=dpi)
        page_count = len(images)

        if page_count == 1:
            # single page → direct save
            out_path = os.path.join(out_dir, f"{base_name}_1.{fmt}")
            images[0].save(out_path, fmt.upper())
            print(f"✅ Single-page PDF converted: {out_path}")
            return out_dir

        # multi-page case → ask user choice
        print("PDF has", page_count, "pages.")
        print("Choose output option:")
        print("1 → Merge all pages into ONE image")
        print("2 → Save each page as separate image")
        choice = input("Select option (1/2): ").strip()

        if choice == "1":
            # merge vertically into single tall image
            widths, heights = zip(*(img.size for img in images))
            total_height = sum(heights)
            max_width = max(widths)

            merged_img = Image.new("RGB", (max_width, total_height), (255, 255, 255))

            y_offset = 0
            for img in images:
                merged_img.paste(img, (0, y_offset))
                y_offset += img.height

            out_path = os.path.join(out_dir, f"{base_name}_merged.{fmt}")
            merged_img.save(out_path, fmt.upper())
            print(f"✅ All pages merged into one image: {out_path}")

        else:
            # separate images
            for i, img in enumerate(images, 1):
                img.save(os.path.join(out_dir, f"{base_name}_{i}.{fmt}"), fmt.upper())
            print(f"✅ {page_count} pages saved as separate images in {out_dir}")

        return out_dir

    except Exception as e:
        print(f"❌ PDF to IMAGE failed: {e}")
        return None

def pdf_to_csv(pdf_path, csv_path=None, xlsx_path=None, csv_delimiter=", " , excel_font_map=FALLBACK_FONTS_PDF, batch_log_every=1000):
    """
    PDF → CSV/XLSX converter (strict mode):
    - Converts only actual tables
    - Ignores all other content (JSON/plain text/XLS)
    - Exits if no tables found
    """
    if not pdf_path.endswith(".pdf"):
        raise RuntimeError("❌ Only PDF files are supported!")

    if not os.path.exists(pdf_path):
        raise RuntimeError("❌ File not found!")

    csv_file = None
    csv_writer = None
    if csv_path:
        csv_file = open(csv_path, "w", encoding="utf-8-sig", newline="")
        csv_writer = csv.writer(csv_file, delimiter=csv_delimiter)

    wb = None
    ws = None
    if xlsx_path:
        wb = Workbook(write_only=True)
        ws = wb.create_sheet(title="Extracted")

    total_rows = 0
    unsupported_pages = []
    valid_table_found = False

    try:
        with pdfplumber.open(pdf_path) as pdf:
            num_pages = len(pdf.pages)
            print(f"Opened PDF: {pdf_path} (pages: {num_pages})")

            for pageno, page in enumerate(pdf.pages, start=1):
                rows = _page_tables_to_rows(page)

                # Strict table check: must have >1 row and >1 column
                if not rows or len(rows) < 1 or all(len(r) < 2 for r in rows):
                    unsupported_pages.append(pageno)
                    continue

                # Valid table found
                valid_table_found = True

                # Write rows
                for row in rows:
                    out_row = [("" if c is None else str(c)) for c in row]

                    if csv_writer:
                        try:
                            csv_writer.writerow(out_row)
                        except Exception:
                            safe_row = [s.encode("utf-8", errors="ignore").decode("utf-8") for s in out_row]
                            csv_writer.writerow(safe_row)

                    if ws:
                        cells = []
                        for cell_val in out_row:
                            script = detect_script_simple(cell_val)
                            font_name = excel_font_map.get(script, excel_font_map.get("DEFAULT", "Arial"))
                            cell = WriteOnlyCell(ws, value=cell_val)
                            cell.font = Font(name=font_name, size=11)
                            cells.append(cell)
                        ws.append(cells)

                    total_rows += 1
                    if total_rows % batch_log_every == 0:
                        print(f"Processed rows: {total_rows} (page {pageno}/{num_pages})")

        if not valid_table_found:
            # Close any opened files
            if csv_file:
                csv_file.close()
                os.remove(csv_path)  # remove empty CSV
            if wb:
                os.remove(xlsx_path)
            print("❌ No embedded table/CSV found in PDF. Only actual tables can be converted in strict mode.")
            return None

        # Close/save files
        if csv_file:
            csv_file.close()
        if wb:
            wb.save(xlsx_path)

        if unsupported_pages:
            print(f"⚠️ Pages skipped (unsupported content): {unsupported_pages}")
            print("Only actual tables converted. JSON/XLS/plain text ignored.")

        print(f"✅ Done. Rows written: {total_rows}")
        out = {}
        if csv_path:
            out["csv"] = os.path.abspath(csv_path)
        if xlsx_path:
            out["xlsx"] = os.path.abspath(xlsx_path)
        out["rows"] = total_rows
        return out

    except Exception as exc:
        try:
            if csv_file and not csv_file.closed:
                csv_file.close()
        except Exception:
            pass
        try:
            if wb:
                wb.save(xlsx_path)
        except Exception:
            pass
        raise RuntimeError(f"Error converting PDF to CSV/XLSX: {exc}") from exc

def pdf_to_xls(pdf_path, xls_path):
    try:
        if not pdf_path.endswith(".pdf"):
            raise RuntimeError("❌ Only PDF files are supported!")

        if not os.path.exists(pdf_path):
            raise RuntimeError("❌ File not found!")

        wb = Workbook(write_only=True)
        ws = wb.create_sheet("Sheet1")
        with pdfplumber.open(pdf_path) as pdf:
            for page in pdf.pages:
                rows = _page_tables_to_rows(page)
                for r in rows:
                    ws.append([str(x) if x is not None else "" for x in r])
        wb.save(xls_path)
        return xls_path

    except Exception as e:
        print(f"❌ PDF to XLS failed: {e}")
        return None

def pdf_to_json(pdf_path, json_path):
    try:
        if not pdf_path.endswith(".pdf"):
            raise RuntimeError("❌ Only PDF files are supported!")

        if not os.path.exists(pdf_path):
            raise RuntimeError("❌ File not found!")

        pages_data = []
        with pdfplumber.open(pdf_path) as pdf:
            for idx, page in enumerate(pdf.pages, 1):
                page_dict = {"page": idx, "text": page.extract_text() or ""}
                # Try tables too
                tables = page.extract_tables() or []
                if tables:
                    page_dict["tables"] = tables
                pages_data.append(page_dict)
        with open(json_path, "w", encoding="utf-8") as out:
            json.dump(pages_data, out, ensure_ascii=False, indent=2)
        return json_path

    except Exception as e:
        print(f"❌ PDF to JSON failed: {e}")
        return None

# =========================
# DOCX Helpers
# =========================

def doc_to_pdf(input_path: str, output_pdf: str, mode: str = "auto", raster_dpi: int = 200) -> str:
    """
    Convert .doc/.docx → .pdf
    - mode="auto":
        * If doc has table(s) with >=5 columns → convert via HTML to PDF (avoids cropping).
        * Else → direct via LibreOffice (fastest, keeps structure).
    - mode="image": LibreOffice → PDF → rasterized (pixel perfect, no font issues).
    """
    try:
        if not os.path.exists(input_path):
            raise RuntimeError("Input file not found: " + input_path)

        _ensure_dir(output_pdf)

        # 1) Ensure DOCX
        docx_path = convert_doc_to_docx_if_needed(input_path)

        # --- Check if doc has wide tables (>=5 columns) ---
        has_wide_tables = False
        try:
            from docx import Document
            doc = Document(docx_path)
            for tbl in doc.tables:
                if tbl.rows and len(tbl.rows[0].cells) >= 5:
                    has_wide_tables = True
                    break
        except Exception as e:
            print(f"⚠️ Could not inspect tables in DOCX: {e}")

        # --- Wide tables? Use HTML → PDF ---
        if has_wide_tables and mode == "auto":
            try:

                with open(docx_path, "rb") as f:
                    html = mammoth.convert_to_html(f).value
                tmp_html = os.path.join(tempfile.gettempdir(), "doc_tmp.html")
                with open(tmp_html, "w", encoding="utf-8") as f:
                    f.write(html)
                pdfkit.from_file(tmp_html, output_pdf)
                print(f"✅ Converted via HTML path (wide table support): {output_pdf}")
                return output_pdf
            except Exception as e:
                print(f"⚠️ HTML conversion failed, falling back to LibreOffice: {e}")

        # --- Default: DOCX → PDF via LibreOffice ---
        tmp_dir = tempfile.mkdtemp(prefix="docx2pdf_")
        lo_pdf = os.path.join(tmp_dir, "lo.pdf")
        try:
            lo_pdf = convert_docx_to_pdf_libreoffice(docx_path, lo_pdf)
        except Exception as e:
            shutil.rmtree(tmp_dir)
            raise RuntimeError(f"LibreOffice conversion failed: {e}")

        if mode == "auto":
            shutil.copy(lo_pdf, output_pdf)
            shutil.rmtree(tmp_dir)
            return output_pdf

        if mode == "image":
            raster_pdf = os.path.join(tmp_dir, "raster.pdf")
            rasterize_pdf_to_pdf(lo_pdf, raster_pdf, dpi=raster_dpi)
            shutil.copy(raster_pdf, output_pdf)
            shutil.rmtree(tmp_dir)
            return output_pdf

        shutil.rmtree(tmp_dir)
        raise ValueError('mode must be "auto" or "image"')

    except Exception as e:
        print(f"❌ DOCX to PDF failed: {e}")
        return None

#def doc_to_txt(docx_path, txt_path):
#    try:
#        ensure_parent_dir(txt_path)
#        doc = Document(docx_path)
#        lines = []
#        for p in doc.paragraphs:
#            lines.append(p.text)
#        with open(txt_path, "w", encoding="utf-8") as out:
#            out.write("\n".join(lines))
#        print(f"✅ DOCX → TXT: {txt_path}")
#        return txt_path
#    except Exception as e:
#        print(f"❌ DOCX to TXT failed: {e}")
#        return None

def doc_to_txt(docx_path, txt_path):
    try:
        if not docx_path.endswith(".docs"):
            raise RuntimeError("❌ Only DOCX files are supported!")

        if not os.path.exists(docx_path):
            raise RuntimeError("❌ File not found!")

        doc = Document(docx_path)
        with open(txt_path, "w", encoding="utf-8") as out:
            # Paragraphs
            for para in doc.paragraphs:
                out.write((para.text or "") + "\n")
            # Tables (append as TSV)
            for table in doc.tables:
                for row in table.rows:
                    cells = [c.text.replace("\n"," ").strip() for c in row.cells]
                    out.write("\t".join(cells) + "\n")
        return txt_path

    except Exception as e:
        print(f"❌ DOCX to TXT failed: {e}")
        return None

def doc_to_docx_image(input_path: str) -> str:
    """
    Ensure the input is in DOCX format.
    If it's a .doc, convert using LibreOffice headless.
    """
    import tempfile, subprocess, os, glob

    if input_path.lower().endswith(".docx"):
        return input_path

    if input_path.lower().endswith(".doc"):
        tmp_dir = tempfile.mkdtemp(prefix="doc2docx_")
        try:
            subprocess.check_call([
                "soffice", "--headless", "--convert-to", "docx",
                "--outdir", tmp_dir, input_path
            ])

            # ✅ Find the .docx file inside tmp_dir
            converted_files = glob.glob(os.path.join(tmp_dir, "*.docx"))
            if not converted_files:
                raise RuntimeError("DOC→DOCX conversion failed, file not created.")

            return converted_files[0]  # Return the actual converted path
        except Exception as e:
            raise RuntimeError(f"Failed to convert .doc to .docx: {e}")
    else:
        raise RuntimeError(f"Unsupported file type: {input_path}")

def doc_to_image(input_path, out_path, dpi=200):
    """
    Convert .doc/.docx → PNG image(s) with full formatting (tables, images, layout preserved).
    - If 1 page → single PNG
    - If multiple pages → user chooses:
        1 = merge into single long PNG
        2 = separate PNG per page (zipped)
    """
    tmp_dir = None
    try:
        # ✅ Always ensure DOCX
        docx_path = doc_to_docx_image(input_path)
        # ✅ Detect wide tables
        wide_table = False
        try:
            doc = Document(docx_path)
            for table in doc.tables:
                if len(table.columns) > 5:
                    wide_table = True
                    break
        except Exception as e:
            print(f"⚠️ Could not inspect tables: {e}")

        # ========== PATH 1: Wide Table → Mammoth + WeasyPrint ==========
        if wide_table:
            print("⚠️ Wide table detected → Using HTML → PDF → Image pipeline")
            tmp_dir = tempfile.mkdtemp(prefix="doc2html_")

            # DOCX → HTML
            with open(docx_path, "rb") as f:
                result = mammoth.convert_to_html(f)
                html_content = result.value

            # Save HTML to file (safe for large tables)
            html_file = os.path.join(tmp_dir, "doc.html")
            with open(html_file, "w", encoding="utf-8") as f:
                f.write(html_content)

            # HTML → PDF
            pdf_path = os.path.join(tmp_dir, "out.pdf")
            pdfkit.from_file(html_file, pdf_path)

        # ========== PATH 2: Normal → LibreOffice ==========
        else:
            print("✅ No wide tables → Using LibreOffice pipeline")
            tmp_dir = tempfile.mkdtemp(prefix="doc2pdf_")
            subprocess.check_call([
                "soffice", "--headless", "--convert-to", "pdf",
                "--outdir", tmp_dir, docx_path
            ])

            pdf_files = glob.glob(os.path.join(tmp_dir, "*.pdf"))
            if not pdf_files:
                raise RuntimeError("DOCX→PDF failed, no PDF generated.")
            pdf_path = pdf_files[0]

        print(f"✅ PDF ready at: {pdf_path}", '3')

        # ✅ PDF → PNG(s)
        pages = convert_from_path(pdf_path, dpi=dpi)
        print(f"pages: {len(pages)}", "4")

        if len(pages) == 1:
            if not out_path.lower().endswith(".png"):
                out_path += ".png"
            pages[0].save(out_path, "PNG")
            print(f"✅ Converted single page image: {out_path}")
            return out_path

        # Multiple pages → ask user
        print("Document has multiple pages. Choose output format:")
        print("1. Single PNG (all pages merged vertically)")
        print("2. Separate PNG per page (zipped)")
        choice = input("Enter choice (1/2): ").strip()

        if choice == "1":
            widths, heights = zip(*(p.size for p in pages))
            total_height = sum(heights)
            max_width = max(widths)
            merged_img = Image.new("RGB", (max_width, total_height), "white")
            y = 0
            for p in pages:
                merged_img.paste(p, (0, y))
                y += p.height
            if not out_path.lower().endswith(".png"):
                out_path += ".png"
            merged_img.save(out_path, "PNG")
            print(f"✅ Merged multi-page image saved: {out_path}")
            return out_path

        elif choice == "2":
            if not out_path.lower().endswith(".zip"):
                out_path += ".zip"
            img_dir = tempfile.mkdtemp(prefix="doc2img_")
            image_files = []
            for i, page in enumerate(pages, 1):
                img_path = os.path.join(img_dir, f"page_{i}.png")
                page.save(img_path, "PNG")
                image_files.append(img_path)

            with zipfile.ZipFile(out_path, "w") as zipf:
                for f in image_files:
                    zipf.write(f, os.path.basename(f))
            print(f"✅ Separate page images saved as zip: {out_path}")
            return out_path

        else:
            raise ValueError("Invalid choice")

    except Exception as e:
        raise RuntimeError(f"Error converting DOCX to PNG: {e}") from e
    finally:
        if tmp_dir and os.path.exists(tmp_dir):
            shutil.rmtree(tmp_dir, ignore_errors=True)

def doc_to_csv(input_path, csv_path):
    try:
        # Step 1: normalize DOC → DOCX
        docx_path = doc_to_docx_image(input_path)
        doc = Document(docx_path)

        # 🚫 Step 2: Reject if any images exist
        for rel in doc.part.rels.values():
            if "image" in rel.target_ref:
                raise RuntimeError("❌ Not supported: Document contains images, cannot export to CSV.")

        # Step 3: Scan all tables first
        valid_tables = []
        for table in doc.tables:
            col_count = len(table.rows[0].cells) if table.rows else 0
            row_count = len(table.rows)

            if col_count < 4 or row_count < 2:
                continue

            meaningful_rows = 0
            for row in table.rows:
                texts = [c.text.strip() for c in row.cells if c.text.strip()]
                if len(texts) >= (col_count // 2) and any(len(t) >= 3 for t in texts):
                    meaningful_rows += 1

            if meaningful_rows >= 2:
                valid_tables.append(table)

        if not valid_tables:
            raise RuntimeError("❌ Not supported: No valid table (≥4 columns, ≥2 rows, with real data).")

        # ✅ Step 4: Now open CSV only if conversion is confirmed
        with open(csv_path, "w", encoding="utf-8-sig", newline="") as out:
            writer = csv.writer(out)
            for table in valid_tables:
                for row in table.rows:
                    row_data = []
                    for c in row.cells:
                        text = c.text.replace("\n", " ").strip()
                        if text:
                            text = text.encode("utf-8", "ignore").decode("utf-8")
                        row_data.append(text)
                    writer.writerow(row_data)

        return csv_path

    except Exception as e:
        # ❌ Ensure partial/empty CSV is deleted if error
        if os.path.exists(csv_path):
            os.remove(csv_path)
        raise RuntimeError(f"Error converting DOC/DOCX to CSV: {e}") from e

def doc_to_xls(input_path, xls_path):
    try:
        # ✅ DOC → DOCX ensure
        docx_path = doc_to_docx_image(input_path)

        # ✅ Load DOCX
        doc = Document(docx_path)

        # ✅ Image check (unsupported for XLS export)
        for shape in doc.inline_shapes:
            raise RuntimeError("❌ Not supported: Document contains images, cannot convert to XLS.")

        wrote_any = False
        valid_table_found = False

        wb = Workbook(write_only=True)
        ws = wb.create_sheet("Sheet1")

        # ✅ Table scan
        for table in doc.tables:
            if not table.rows or len(table.rows[0].cells) < 4:
                continue  # skip invalid
            if len(table.rows) < 2:
                continue  # need at least 2 rows

            # check real data
            has_data = any(any(c.text.strip() for c in row.cells) for row in table.rows)
            if not has_data:
                continue

            valid_table_found = True
            for row in table.rows:
                row_data = []
                for c in row.cells:
                    text = c.text.replace("\n", " ").strip()
                    if text:
                        text = text.encode("utf-8", "ignore").decode("utf-8")
                    row_data.append(text)
                ws.append(row_data)
                wrote_any = True

        # ❌ No valid table
        if not valid_table_found:
            raise RuntimeError("❌ Not supported: No valid table (≥4 columns, ≥2 rows, with real data).")

        if not wrote_any:
            raise RuntimeError("❌ Table found but empty, cannot convert.")

        wb.save(xls_path)
        return xls_path

    except Exception as e:
        # ⚠️ Agar error ho to partial file delete
        if os.path.exists(xls_path):
            os.remove(xls_path)
        raise RuntimeError(f"Error converting DOC/DOCX to XLSX: {e}") from e

def doc_to_json(input_path, json_path):
    try:
        # ✅ DOC → DOCX ensure
        docx_path = doc_to_docx_image(input_path)

        # ✅ Load DOCX
        doc = Document(docx_path)

        # ✅ Step 3: Image check
        for rel in doc.part.rels.values():
            if "image" in rel.reltype:
                raise RuntimeError("❌ Not supported: Document contains images.")

        # ✅ Step 4: Raw text join
        raw_text = "\n".join([p.text.strip() for p in doc.paragraphs if p.text.strip()])

        # ✅ Step 5: Detect JSON-like content
        if raw_text.strip().startswith(("{", "[")):
            try:
                parsed = json.loads(raw_text)
                with open(json_path, "w", encoding="utf-8") as out:
                    json.dump(parsed, out, ensure_ascii=False, indent=2)
                return json_path
            except Exception:
                pass  # not valid JSON, fallback

        # ✅ Step 6: Detect table data
        valid_tables = []
        for table in doc.tables:
            if len(table.columns) >= 4 and len(table.rows) >= 2:
                table_data = []
                for row in table.rows:
                    row_data = []
                    for c in row.cells:
                        text = c.text.replace("\n", " ").strip()
                        if text:
                            text = text.encode("utf-8", "ignore").decode("utf-8")
                        row_data.append(text)
                    table_data.append(row_data)
                valid_tables.append(table_data)

        # ✅ Step 7: Detect tab/comma separated plain text (like csv.docx)
        if not valid_tables:
            lines = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
            table_data = []
            for line in lines:
                if "," in line or "\t" in line or ";" in line:
                    parts = re.split(r'[\t,;]', line)
                    if len(parts) >= 4:
                        table_data.append([p.strip() for p in parts])
            if table_data:
                valid_tables.append(table_data)

        # ✅ Step 8: Final check
        if not valid_tables:
            raise RuntimeError("❌ Not supported: No valid JSON or tabular data found.")

        payload = {"tables": valid_tables}
        with open(json_path, "w", encoding="utf-8") as out:
            json.dump(payload, out, ensure_ascii=False, indent=2)

        return json_path

    except Exception as e:
        # ⚠️ Error aane par partial file delete
        if os.path.exists(json_path):
            os.remove(json_path)
        raise RuntimeError(f"Error converting DOC/DOCX to JSON: {e}") from e

# =========================
# Image Helpers
# =========================

def image_to_txt_ocr(in_path, txt_path, lang="auto"):
    """
    Extracts text from an image using OCR. Supports multiple languages.
    Uses EasyOCR and Tesseract, prioritizing EasyOCR for some languages.

    Args:
        in_path (str): Path to the input image file.
        txt_path (str): Path to the output text file.
        lang (str): Language code(s) (e.g., 'en', 'hi', 'en,hi'). Use 'auto' for auto-detection.

    Returns:
        str: Path to the output text file if successful, None otherwise.
    """
    try:
        if not os.path.exists(in_path):
            raise FileNotFoundError(f"❌ File not found: {in_path}")

        # Preprocess image
        img = preprocess_image(in_path, strong=False)

        # Handle 'auto' language detection
        target_langs = []
        if lang.lower() == 'auto':
            # For auto-detection, load default EasyOCR languages and then try to detect
            target_langs_for_easyocr = _parse_langs_for_easyocr(lang)
            reader = _get_easyocr(target_langs_for_easyocr) # Load initial languages
            if reader:
                # Now try to detect language from the image
                detected = detect_language_from_image(img)
                print(f"🌐 Detected language: {detected}")
                # Re-initialize EasyOCR with the detected language + English for robustness
                target_langs = list(set([detected, 'en'])) # Use set to avoid duplicates
                # Ensure the detected language is in the supported list for EasyOCR if not English
                if detected != 'en' and detected not in ['hi', 'es', 'fr', 'de', 'ru', 'ja', 'ko', 'ch_sim', 'ch_tra']: # Add more supported languages if needed
                    print(f"⚠️ Detected language {detected} might not be fully supported by EasyOCR. Using English.")
                    target_langs = ['en'] # Fallback to English
                reader = _get_easyocr(target_langs) # Re-initialize with detected/fallback languages
            else:
                # If EasyOCR couldn't even initialize with defaults
                target_langs = ['en'] # Fallback to Tesseract with English
        else:
            target_langs = _parse_langs_for_easyocr(lang)
            reader = _get_easyocr(target_langs) # Initialize with specified languages

        text = ""
        used_engine = ""

        # --- Attempt with EasyOCR ---
        # Ensure reader is not None before using
        if reader:
            try:
                # Use paragraph=True for better formatting
                results = reader.readtext(img, detail=0, paragraph=True)
                text = "\n".join(results).strip()
                used_engine = "EasyOCR"
                print(f"✅ OCR extracted text using EasyOCR.")
            except Exception as e:
                print(f"⚠️ EasyOCR failed: {e}")


        # --- Fallback or additional pass with Tesseract ---
        if not text or used_engine != "EasyOCR": # If EasyOCR failed or was skipped
            tess_langs = _langs_for_tesseract(target_langs)
            print(f"Attempting Tesseract with languages: {tess_langs}")
            try:
                # Ensure tesseract_cmd is set if not in PATH (common in Colab)
                # pytesseract.tesseract_cmd = r'/usr/bin/tesseract' # Uncomment if needed
                text = pytesseract.image_to_string(img, lang=tess_langs, config="--oem 3 --psm 6").strip()
                used_engine = "Tesseract"
                print(f"✅ OCR extracted text using Tesseract.")
            except pytesseract.TesseractNotFoundError:
                print("❌ Tesseract is not installed or not in PATH. Please install it.")
                text = "" # Ensure text is empty if Tesseract is not found
            except Exception as e:
                print(f"⚠️ Tesseract failed for languages {tess_langs}: {e}")
                text = "" # Ensure text is empty if Tesseract also fails

        if not text:
            raise RuntimeError("OCR produced empty text from both engines or both failed.")

        with open(txt_path, "w", encoding="utf-8") as f:
            f.write(text)
        print(f"✅ OCR extracted text saved: {txt_path} (Engine: {used_engine})")
        return txt_path

    except FileNotFoundError as e:
        print(f"❌ File not found: {e}")
        return None
    except Exception as e:
        print(f"❌ OCR failed: {e}")
        return None

def image_to_image(in_path, out_path, fmt):
    try:
        # fmt: "PNG", "JPEG", etc.
        img = Image.open(in_path)
        # Convert mode if needed for JPEG
        if fmt.upper() in ["JPEG", "JPG"] and img.mode in ("RGBA", "LA"):
            bg = Image.new("RGB", img.size, (255,255,255))
            bg.paste(img, mask=img.split()[-1])
            img = bg
        img.convert("RGB").save(out_path, fmt.upper())
        return out_path

    except Exception as e:
        raise RuntimeError(f"❌ Image conversion failed: {e}")


# =========================
# Translation
# =========================

def translate_text(text, src_lang="auto", dest_lang="en"):
    translated_chunks = []
    chunk_size = 4000  # Google limit-ish
    for i in range(0, len(text), chunk_size):
        chunk = text[i:i+chunk_size]
        translated = GoogleTranslator(source=src_lang, target=dest_lang).translate(chunk)
        translated_chunks.append(translated)
    return "\n".join(translated_chunks)

def translate_file(file_path, out_path, src_lang="auto", dest_lang="en"):
    try:
        ensure_parent_dir(out_path)
        ext = os.path.splitext(file_path)[1].lower()
        if ext == ".pdf":
            temp_txt = os.path.splitext(out_path)[0] + ".__tmp_in.txt"
            txt_extracted = pdf_to_txt(file_path, temp_txt)
            if not txt_extracted: return None
            with open(temp_txt, "r", encoding="utf-8") as f:
                content = f.read()
            try:
                os.remove(temp_txt)
            except Exception:
                pass
        elif ext == ".docx":
            temp_txt = os.path.splitext(out_path)[0] + ".__tmp_in.txt"
            txt_extracted = doc_to_txt(file_path, temp_txt)
            if not txt_extracted: return None
            with open(temp_txt, "r", encoding="utf-8") as f:
                content = f.read()
            try:
                os.remove(temp_txt)
            except Exception:
                pass
        else:
            with open(file_path, "r", encoding="utf-8") as f:
                content = f.read()

        translated_text = translate_text(content, src_lang, dest_lang)
        with open(out_path, "w", encoding="utf-8") as f:
            f.write(translated_text)
        print(f"✅ Translated → {out_path}")
        return out_path
    except Exception as e:
        print(f"❌ Translation failed: {e}")
        return None

# ---- Tiny wrappers (unchanged names) ----
def jpg_to_png(in_path, out_path): return image_to_image(in_path, out_path, "PNG")
def png_to_jpg(in_path, out_path): return image_to_image(in_path, out_path, "JPEG")
def jpg_to_jpeg(in_path, out_path): return image_to_image(in_path, out_path, "JPEG")
def png_to_jpeg(in_path, out_path): return image_to_image(in_path, out_path, "JPEG")
def jpeg_to_png(in_path, out_path): return image_to_image(in_path, out_path, "PNG")
def jpeg_to_jpg(in_path, out_path): return image_to_image(in_path, out_path, "JPEG")
def gif_to_png(in_path, out_path): return image_to_image(in_path, out_path, "PNG")

# OCR wrappers (note: lang now supports "auto" / "en,hi" / "eng+hin")
def jpg_to_txt(in_path, out_path, lang="auto"):  return image_to_txt_ocr(in_path, out_path, lang)
def jpeg_to_txt(in_path, out_path, lang="auto"): return image_to_txt_ocr(in_path, out_path, lang)
def png_to_txt(in_path, out_path, lang="auto"):  return image_to_txt_ocr(in_path, out_path, lang)
def gif_to_txt(in_path, out_path, lang="auto"):  return image_to_txt_ocr(in_path, out_path, lang)
def tiff_to_txt(in_path, out_path, lang="auto"): return image_to_txt_ocr(in_path, out_path, lang)
def bmp_to_txt(in_path, out_path, lang="auto"):  return image_to_txt_ocr(in_path, out_path, lang)

# =========================
# Dispatcher
# =========================

CONVERTERS = {
    # CSV →
    ("csv", "xlsx"): csv_to_xls,
    ("csv", "pdf"): csv_to_pdf,
    ("csv", "docx"): csv_to_doc,
    ("csv", "txt"): csv_to_txt,
    ("csv", "json"): csv_to_json,
    ("csv", "png"): csv_to_image,

    # XLSX →
    ("xlsx", "csv"): xls_to_csv,
    ("xlsx", "docx"): xls_to_doc,
    ("xlsx", "txt"): xls_to_txt,
    ("xlsx", "pdf"): xls_to_pdf,
    ("xlsx", "json"): xls_to_json,
    ("xlsx", "png"): xls_to_image,

    # TXT →
    ("txt", "pdf"): txt_to_pdf,
    ("txt", "docx"): txt_to_doc,
    ("txt", "json"): txt_to_json,
    ("txt", "csv"): txt_to_csv,
    ("txt", "png"): txt_to_image,
    ("txt", "xlsx"): txt_to_xls,

    # JSON →
    ("json", "csv"): json_to_csv,
    ("json", "xlsx"): json_to_xls,
    ("json", "txt"): json_to_txt,
    ("json", "pdf"): json_to_pdf,
    ("json", "docx"): json_to_doc,
    ("json", "png"): json_to_image,

    # DOCX → (helpers)
    ("docx", "txt"): doc_to_txt,
    ("docx", "pdf"): doc_to_pdf,
    ("docx", "xls"): doc_to_xls,
    ("docx", "csv"): doc_to_csv,
    ("docx", "json"): doc_to_json,
    ("docx", "image"): doc_to_image,

    # PDF → (helpers)
    ("pdf", "txt"): pdf_to_txt,
    ("pdf", "docx"): pdf_to_docx,
    ("pdf", "image"): pdf_to_image,
    ("pdf", "csv"): pdf_to_csv,
    ("pdf", "xls"): pdf_to_xls,
    ("pdf", "json"): pdf_to_json,
    ("pdf", "txt ocr"): pdf_to_txt_ocr,

    # IMAGE → (helpers)
    ("png", "jpg"): image_to_image,
    ("png", "jpeg"): image_to_image,
    ("jpg", "jpeg"): image_to_image,
    ("jpg", "png"): image_to_image,
    ("jpeg", "jpg"): image_to_image,
    ("jpeg", "png"): image_to_image,
    ("gif", "png") : image_to_image,

    ("jpg", "txt"): image_to_txt_ocr,
    ("jpeg", "txt"): image_to_txt_ocr,
    ("png", "txt"): image_to_txt_ocr,
    ("gif", "txt"): image_to_txt_ocr,
    ("tiff", "txt"): image_to_txt_ocr,
    ("bmp", "txt"): image_to_txt_ocr,

}

SUPPORTED_MENUS = {
    "pdf": [
        ("1", "PDF → TXT", "txt"),
        ("2", "PDF → DOCX", "docx"),
        ("3", "PDF → PNG (image)", "png"),
        ("4", "PDF → CSV", "csv"),
        ("5", "PDF → XLSX", "xlsx"),
        ("6", "PDF → JSON", "json"),
        ("7", "PDF → TXT OCR", "txt ocr"),
    ],
    "csv": [
        ("1", "CSV → PDF", "pdf"),
        ("2", "CSV → TXT", "txt"),
        ("3", "CSV → PNG (image)", "png"),
        ("4", "CSV → DOCX", "docx"),
        ("5", "CSV → XLSX", "xlsx"),
        ("6", "CSV → JSON", "json"),
    ],
    "xls": [
        ("1", "XLSX → CSV", "csv"),
        ("2", "XLSX → PDF", "pdf"),
        ("3", "XLSX → TXT", "txt"),
        ("4", "XLSX → DOCX", "docx"),
        ("5", "XLSX → JSON", "json"),
        ("6", "XLSX → PNG (image)", "png"),
    ],
    "txt": [
        ("1", "TXT → CSV", "csv"),
        ("2", "TXT → PDF", "pdf"),
        ("3", "TXT → DOCX", "docx"),
        ("4", "TXT → JSON", "json"),
        ("5", "TXT → XLSX", "xlsx"),
        ("6", "TXT → PNG (image)", "png"),
    ],
    "json": [
        ("1", "JSON → CSV", "csv"),
        ("2", "JSON → TXT", "txt"),
        ("3", "JSON → XLSX", "xls"),
        ("4", "JSON → DOCX", "docx"),
        ("5", "JSON → PDF", "pdf"),
        ("6", "JSON → PNG (image)", "png"),
    ],
    "image": [
        ("1", "JPG → PNG", "png"),
        ("2", "PNG → JPG", "jpg"),
        ("3", "JPG → JPEG", "jpeg"),
        ("4", "PNG → JPEG", "jpeg"),
        ("5", "JPEG → PNG", "png"),
        ("6", "JPEG → JPG", "jpg"),
        ("7", "GIF → PNG", "png"),
    ],
}

def infer_ext(path):
    return os.path.splitext(path)[1].lower().strip(".") if "." in os.path.basename(path) else ""

def run_conversion(src_fmt, dst_fmt, in_path, out_path):
    key = (src_fmt.lower(), dst_fmt.lower())
    if key not in CONVERTERS:
        print(f"❌ Conversion not supported: {src_fmt} → {dst_fmt}")
        return None
    return CONVERTERS[key](in_path, out_path)

# =========================
# CLI Menu
# =========================

def cli_menu():

    try:
        from google.colab import files
    except ImportError:
        files = None

    while True:
        print("\n===== File Toolkit =====")
        print("1. Convert File")
        print("2. Translate File")
        print("3. Exit")
        choice = input("Enter choice: ").strip()

        if choice == "1":
            try:
                # Upload file (Colab) or enter path (offline)
                if files:
                    uploaded = files.upload()
                    in_path = list(uploaded.keys())[0]
                else:
                    in_path = input("Enter path of File): ").strip()

                #in_path = input("Enter input file path: ").strip()
                if not os.path.exists(in_path):
                    print("❌ Input file not found.")
                    continue

                src_fmt = input("Enter source format (auto to detect): ").strip().lower()
                if src_fmt in ("", "auto"):
                    src_fmt = infer_ext(in_path)
                    if not src_fmt:
                        print("❌ Could not infer source format. Provide explicitly.")
                        continue

                # ✅ NEW DYNAMIC MENU
                if src_fmt in SUPPORTED_MENUS:
                    print("\nSupported conversions:")
                    for opt, desc, _ in SUPPORTED_MENUS[src_fmt]:
                        print(f"{opt}. {desc}")
                    fmt_choice = input("Select target format: ").strip()

                    match = next((m for m in SUPPORTED_MENUS[src_fmt] if m[0] == fmt_choice), None)
                    if not match:
                        print("❌ Invalid choice!")
                        continue

                    _, _, dst_fmt = match
                else:
                    dst_fmt = input("Enter target format: ").strip().lower()
                    if not dst_fmt:
                        print("❌ Target format required.")
                        continue

                # Output path
                out_path = input("Enter output file path (leave blank to auto): ").strip()
                if not out_path:
                    base, _ = os.path.splitext(in_path)
                    out_path = f"{base}.{dst_fmt}"

                # Run conversion
                func = CONVERTERS.get((src_fmt, dst_fmt))
                if func:
                    try:
                        result = func(in_path, out_path)
                        if result:
                            print(f"✅ Converted successfully: {result}")
                    except Exception as e:
                        print(f"❌ Conversion failed: {e}")
                else:
                    print(f"❌ Conversion not supported: {src_fmt} → {dst_fmt}")

            except Exception as e:
                print(f"❌ Error: {e}")

        elif choice == "2":
            try:
                in_path = input("Enter file path to translate: ").strip()
                if not os.path.exists(in_path):
                    print("❌ Input file not found.")
                    continue
                src_lang = input("Source language (default 'auto'): ").strip() or "auto"
                dst_lang = input("Target language (default 'en'): ").strip() or "en"
                out_path = input("Output text file path (e.g., translated.txt): ").strip()
                if not out_path:
                    base, _ = os.path.splitext(in_path)
                    out_path = f"{base}.translated.txt"

                result = translate_file(in_path, out_path, src_lang=src_lang, dest_lang=dst_lang)
                if result:
                    print(f"✅ Translated successfully: {result}")
                    # schedule_delete(result)
                else:
                    print("❌ Translation failed.")
            except Exception as e:
                print(f"❌ Error: {e}")

        elif choice == "3":
            print("👋 Exiting...")
            break
        else:
            print("❌ Invalid choice!")

# =========================
# Entry
# =========================

if __name__ == "__main__":
    # If you also want a non-interactive mode, you can add argparse here.
    # For now, we use the requested menu UI:
    cli_menu()
