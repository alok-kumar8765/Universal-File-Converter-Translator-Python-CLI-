# pdf_converters.py
import os, csv, json
import pdfplumber
from docx import Document
from pdf2image import convert_from_path
import pandas as pd
from openpyxl import Workbook
from docx.shared import Pt
from docx.shared import Inches
from docx.oxml.ns import qn
from pdf2image import convert_from_path
from PIL import Image
import unicodedata
from openpyxl.styles import Font

def pdf_to_txt(pdf_path, txt_path):
    with pdfplumber.open(pdf_path) as pdf, open(txt_path, "w", encoding="utf-8") as out:
        for page in pdf.pages:
            text = page.extract_text() or ""
            out.write(text.rstrip() + "\n\n")
    return txt_path


# ---- Script to Font mapping ----
FALLBACK_FONTS = {
    "LATIN": "Times New Roman",
    "DEVANAGARI": "Nirmala UI",        # Hindi
    "CJK": "SimSun",                   # Chinese/Japanese/Korean
    "ARABIC": "Amiri",
    "GREEK": "Palatino Linotype",
    "DEFAULT": "Arial"
}

def detect_script(text):
    """Detects script of the text based on Unicode names"""
    for ch in text:
        name = unicodedata.name(ch, "")
        if "DEVANAGARI" in name:
            return "DEVANAGARI"
        elif any(x in name for x in ["CJK UNIFIED", "HIRAGANA", "KATAKANA", "HANGUL"]):
            return "CJK"
        elif "ARABIC" in name:
            return "ARABIC"
        elif "GREEK" in name:
            return "GREEK"
    return "LATIN"

def pdf_to_doc(pdf_path, docx_path):
    try:
        doc = Document()

        with pdfplumber.open(pdf_path) as pdf:
            for pageno, page in enumerate(pdf.pages, 1):

                # --- Extract Text ---
                text = page.extract_text() or ""
                for line in text.splitlines():
                    script = detect_script(line)
                    font_name = FALLBACK_FONTS.get(script, FALLBACK_FONTS["DEFAULT"])

                    para = doc.add_paragraph()
                    run = para.add_run(line)
                    run.font.name = font_name
                    run.font.size = Pt(11)

                    # set for East Asian / Complex Scripts
                    rPr = run._element.rPr.rFonts
                    rPr.set(qn('w:eastAsia'), font_name)
                    rPr.set(qn('w:cs'), font_name)
                    rPr.set(qn('w:hAnsi'), font_name)

                # --- Extract Images ---
                for img in page.images:
                    try:
                        x0, top, x1, bottom = img["x0"], img["top"], img["x1"], img["bottom"]
                        cropped = page.crop((x0, top, x1, bottom))
                        pil_img = cropped.to_image(resolution=150).original
                        img_path = f"temp_img_{pageno}.png"
                        pil_img.save(img_path)

                        # insert into docx
                        doc.add_picture(img_path, width=Inches(4))
                    except Exception as e:
                        print(f"⚠️ Image on page {pageno} skipped: {e}")

                if pageno < len(pdf.pages):
                    doc.add_page_break()

        doc.save(docx_path)
        return docx_path

    except Exception as e:
        print(f"⚠️ Error: {e}")
        return None

def pdf_to_image(pdf_path, out_dir, fmt="png", dpi=150, base_name="page"):
    try:
      os.makedirs(out_dir, exist_ok=True)

      # convert all pages
      images = convert_from_path(pdf_path, dpi=dpi)
      page_count = len(images)

      if page_count == 1:
          # single page → direct save
          out_path = os.path.join(out_dir, f"{base_name}_1.{fmt}")
          images[0].save(out_path, fmt.upper())
          print(f"✅ Single-page PDF converted: {out_path}")
          return out_dir

      # multi-page case → ask user choice
      print("PDF has", page_count, "pages.")
      print("Choose output option:")
      print("1 → Merge all pages into ONE image")
      print("2 → Save each page as separate image")
      choice = input("Select option (1/2): ").strip()

      if choice == "1":
          # merge vertically into single tall image
          widths, heights = zip(*(img.size for img in images))
          total_height = sum(heights)
          max_width = max(widths)

          merged_img = Image.new("RGB", (max_width, total_height), (255, 255, 255))

          y_offset = 0
          for img in images:
              merged_img.paste(img, (0, y_offset))
              y_offset += img.height

          out_path = os.path.join(out_dir, f"{base_name}_merged.{fmt}")
          merged_img.save(out_path, fmt.upper())
          print(f"✅ All pages merged into one image: {out_path}")

      else:
          # separate images
          for i, img in enumerate(images, 1):
              img.save(os.path.join(out_dir, f"{base_name}_{i}.{fmt}"), fmt.upper())
          print(f"✅ {page_count} pages saved as separate images in {out_dir}")

      return out_dir

    except Exception as e:
        print(f"⚠️ Error: {e}")
        return None

# ---------- Helpers: script detection & font mapping ----------
def detect_script_simple(text):
    """
    Heuristic script detection based on Unicode codepoints.
    Returns one of: DEVANAGARI, CJK, ARABIC, GREEK, HANGUL, LATIN, OTHER
    """
    if not text:
        return "LATIN"
    for ch in text:
        cp = ord(ch)
        # Devanagari
        if 0x0900 <= cp <= 0x097F:
            return "DEVANAGARI"
        # CJK unified ideographs (Chinese)
        if 0x4E00 <= cp <= 0x9FFF or 0x3400 <= cp <= 0x4DBF:
            return "CJK"
        # Hiragana / Katakana / Kanji ranges also considered CJK/Japanese
        if 0x3040 <= cp <= 0x30FF:
            return "CJK"
        # Hangul (Korean)
        if 0xAC00 <= cp <= 0xD7AF:
            return "HANGUL"
        # Arabic
        if 0x0600 <= cp <= 0x06FF or 0x0750 <= cp <= 0x077F:
            return "ARABIC"
        # Greek
        if 0x0370 <= cp <= 0x03FF:
            return "GREEK"
        # Basic Latin (ASCII)
        if cp <= 0x007F:
            continue
        # If many other non-ascii chars, mark OTHER as fallback
    # fallback: if there are non-ascii chars but not matched above
    if any(ord(ch) > 127 for ch in text):
        return "OTHER"
    return "LATIN"

# Font mapping for XLSX cells (set to fonts commonly available on Windows; adjust if you prefer others)
FONT_MAP = {
    "DEVANAGARI": "Nirmala UI",         # good for Hindi on Windows
    "CJK": "Microsoft YaHei",           # Simplified Chinese / good on Windows
    "HANGUL": "Malgun Gothic",          # Korean on Windows
    "ARABIC": "Scheherazade",           # try Arabic-friendly; fallback may occur
    "GREEK": "Palatino Linotype",
    "OTHER": "Arial Unicode MS",        # broad coverage (if present)
    "LATIN": "Arial"
}

def _page_tables_to_rows(page):
    """
    Try to extract tables from a pdfplumber page as rows.
    If no tables detected, fall back to extracting text lines and splitting heuristically.
    Returns a list of row-lists.
    """
    rows = []
    # 1) Try extract_tables (pdfplumber)
    try:
        tables = page.extract_tables()  # returns list of tables, each table is list of rows
        if tables:
            for table in tables:
                # table: list of rows (cells may be None)
                for row in table:
                    rows.append([cell for cell in row])
            if rows:
                return rows
    except Exception:
        # continue to fallback
        pass

    # 2) Fallback: use page.extract_text() and split lines by common delimiters (tab or comma)
    text = page.extract_text() or ""
    if not text:
        return []  # nothing to do
    # attempt to detect delimiter by checking first non-empty line
    lines = [ln for ln in text.splitlines() if ln.strip()]
    if not lines:
        return []
    first = lines[0]
    # choose delimiter heuristically
    delim = "\t"
    if "," in first and first.count(",") >= 1:
        delim = ","
    elif "|" in first and first.count("|") >= 1:
        delim = "|"
    # split each line
    for ln in lines:
        parts = [p.strip() for p in ln.split(delim)]
        rows.append(parts)
    return rows

# ---------- Main function: writes CSV (streaming) and XLSX (write_only) ----------
def pdf_to_csv(pdf_path, csv_path=None, xlsx_path=None,
                      csv_delimiter=",", excel_font_map=FONT_MAP,
                      batch_log_every=1000):
    """
    PDF → CSV/XLSX converter (strict mode):
    - Converts only actual tables
    - Ignores all other content (JSON/plain text/XLS)
    - Exits if no tables found
    """
    if not csv_path and not xlsx_path:
        raise ValueError("Provide at least one of csv_path or xlsx_path")

    csv_file = None
    csv_writer = None
    if csv_path:
        csv_file = open(csv_path, "w", encoding="utf-8-sig", newline="")
        csv_writer = csv.writer(csv_file, delimiter=csv_delimiter)

    wb = None
    ws = None
    if xlsx_path:
        wb = Workbook(write_only=True)
        ws = wb.create_sheet(title="Extracted")

    total_rows = 0
    unsupported_pages = []
    valid_table_found = False

    try:
        with pdfplumber.open(pdf_path) as pdf:
            num_pages = len(pdf.pages)
            print(f"Opened PDF: {pdf_path} (pages: {num_pages})")

            for pageno, page in enumerate(pdf.pages, start=1):
                rows = _page_tables_to_rows(page)

                # Strict table check: must have >1 row and >1 column
                if not rows or len(rows) < 1 or all(len(r) < 2 for r in rows):
                    unsupported_pages.append(pageno)
                    continue

                # Valid table found
                valid_table_found = True

                # Write rows
                for row in rows:
                    out_row = [("" if c is None else str(c)) for c in row]

                    if csv_writer:
                        try:
                            csv_writer.writerow(out_row)
                        except Exception:
                            safe_row = [s.encode("utf-8", errors="ignore").decode("utf-8") for s in out_row]
                            csv_writer.writerow(safe_row)

                    if ws:
                        cells = []
                        for cell_val in out_row:
                            script = detect_script_simple(cell_val)
                            font_name = excel_font_map.get(script, excel_font_map.get("DEFAULT", "Arial"))
                            cell = WriteOnlyCell(ws, value=cell_val)
                            cell.font = Font(name=font_name, size=11)
                            cells.append(cell)
                        ws.append(cells)

                    total_rows += 1
                    if total_rows % batch_log_every == 0:
                        print(f"Processed rows: {total_rows} (page {pageno}/{num_pages})")

        if not valid_table_found:
            # Close any opened files
            if csv_file:
                csv_file.close()
                os.remove(csv_path)  # remove empty CSV
            if wb:
                os.remove(xlsx_path)
            print("❌ No embedded table/CSV found in PDF. Only actual tables can be converted in strict mode.")
            return None

        # Close/save files
        if csv_file:
            csv_file.close()
        if wb:
            wb.save(xlsx_path)

        if unsupported_pages:
            print(f"⚠️ Pages skipped (unsupported content): {unsupported_pages}")
            print("Only actual tables converted. JSON/XLS/plain text ignored.")

        print(f"✅ Done. Rows written: {total_rows}")
        out = {}
        if csv_path:
            out["csv"] = os.path.abspath(csv_path)
        if xlsx_path:
            out["xlsx"] = os.path.abspath(xlsx_path)
        out["rows"] = total_rows
        return out

    except Exception as exc:
        try:
            if csv_file and not csv_file.closed:
                csv_file.close()
        except Exception:
            pass
        try:
            if wb:
                wb.save(xlsx_path)
        except Exception:
            pass
        raise RuntimeError(f"Error converting PDF to CSV/XLSX: {exc}") from exc



def pdf_to_xls(pdf_path, xls_path):
    wb = Workbook(write_only=True)
    ws = wb.create_sheet("Sheet1")
    with pdfplumber.open(pdf_path) as pdf:
        for page in pdf.pages:
            rows = _page_tables_to_rows(page)
            for r in rows:
                ws.append([str(x) if x is not None else "" for x in r])
    wb.save(xls_path)
    return xls_path

def pdf_to_json(pdf_path, json_path):
    pages_data = []
    with pdfplumber.open(pdf_path) as pdf:
        for idx, page in enumerate(pdf.pages, 1):
            page_dict = {"page": idx, "text": page.extract_text() or ""}
            # Try tables too
            tables = page.extract_tables() or []
            if tables:
                page_dict["tables"] = tables
            pages_data.append(page_dict)
    with open(json_path, "w", encoding="utf-8") as out:
        json.dump(pages_data, out, ensure_ascii=False, indent=2)
    return json_path

# ---------------- CLI MENU ---------------- #

def main():
    # For Colab file handling
    try:
        from google.colab import files
    except ImportError:
        files = None

    while True:
        print("\n==== PDF File Converter =====")
        print("1. Convert PDF File")
        print("2. Exit")
        choice = input("Enter choice: ").strip()

        if choice == "1":
            # Upload file (Colab) or enter path (offline)
            if files:
                #uploaded = files.upload()
                #pdf_file = list(uploaded.keys())[0]
                pdf_file = input("Enter path of PDF file: ").strip()
            else:
                pdf_file = input("Enter path of PDF file: ").strip()

            print("\nSupported conversions: ")
            print("1. PDF → TXT")
            print("2. PDF → DOCX")
            print("3. PDF → PNG (image)")
            print("4. PDF → CSV")
            print("5. PDF → XLSX")
            print("6. PDF → JSON")
            fmt_choice = input("Select target format (1-6): ").strip()

            base, _ = os.path.splitext(pdf_file)
            out_file = None

            try:
                if fmt_choice == "1":
                    out_file = base + ".txt"
                    pdf_to_txt(pdf_file, out_file)
                elif fmt_choice == "2":
                    out_file = base + ".docx"
                    pdf_to_doc(pdf_file, out_file)
                elif fmt_choice == "3":
                    out_file = base + ".png"
                    out_dir = base + "_images"
                    pdf_to_image(pdf_file, out_dir)
                    out_file = out_dir # set out_file to directory for download
                elif fmt_choice == "4":
                    out_file = base + ".csv"
                    pdf_to_csv(pdf_file, out_file)
                elif fmt_choice == "5":
                    out_file = base + ".xlsx"
                    pdf_to_xls(pdf_file, out_file)
                elif fmt_choice == "6":
                    out_file = base + ".json"
                    pdf_to_json(pdf_file, out_file)
                else:
                    print("❌ Invalid choice!")
                    continue

                print(f"✅ Converted successfully: {out_file}")

                # Colab download option
                if files and out_file:
                    if os.path.isdir(out_file): # if output is a directory (images)
                         print(f"💡 Multiple files saved in {out_file}. You may need to zip and download manually.")
                         # Example of zipping and offering download (requires zip installed)
                         # !zip -r {out_file}.zip {out_file}
                         # files.download(f"{out_file}.zip")
                    else:
                         files.download(out_file)


            except Exception as e:
                print("❌ Conversion failed:", e)

        elif choice == "2":
            print("👋 Exiting...")
            break
        else:
            print("❌ Invalid choice!")

if __name__ == "__main__":
    main()