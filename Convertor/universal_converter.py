#!/usr/bin/env python3
# unified_converter.py

import os
import sys
import csv
import json
import argparse
import threading
import textwrap
import zipfile
import shutil
from io import StringIO

# Third-party libs
import pandas as pd
from openpyxl import Workbook, load_workbook
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import A4
from docx import Document
from pdfminer.high_level import extract_text as pdf_extract_text
from PIL import Image, ImageDraw, ImageFont
from deep_translator import GoogleTranslator

PAGE_W, PAGE_H = A4

# =========================
# Utilities
# =========================

def schedule_delete(file_path, delay=300):
    """Auto-delete a file/folder after delay seconds."""
    try:
        def delete_file():
            if os.path.exists(file_path):
                try:
                    if os.path.isdir(file_path):
                        shutil.rmtree(file_path)
                    else:
                        os.remove(file_path)
                    print(f"🗑 Deleted: {file_path}")
                except Exception as e:
                    print(f"❌ Failed to delete {file_path}: {e}")
        timer = threading.Timer(delay, delete_file)
        timer.start()
    except Exception as e:
        print(f"⚠️ schedule_delete failed: {e}")

def ensure_parent_dir(path: str):
    try:
        parent = os.path.dirname(os.path.abspath(path))
        if parent and not os.path.exists(parent):
            os.makedirs(parent, exist_ok=True)
    except Exception as e:
        print(f"⚠️ Could not ensure parent dir for {path}: {e}")

# ---------- Script to Font mapping (for image conversions) ----------
FALLBACK_FONTS = {
    "LATIN": "arial.ttf",
    "DEVANAGARI": "NotoSansDevanagari-Medium.ttf",
    "CJK": "arialuni.ttf",
    "ARABIC": "arial.ttf",
    "GREEK": "arial.ttf",
    "OTHER": "arialuni.ttf",
    "DEFAULT": "arial.ttf"
}

def detect_script_simple(text: str) -> str:
    if not text:
        return "LATIN"
    for ch in text:
        cp = ord(ch)
        if 0x0900 <= cp <= 0x097F:
            return "DEVANAGARI"
        if (0x4E00 <= cp <= 0x9FFF) or (0x3400 <= cp <= 0x4DBF) or (0x3040 <= cp <= 0x30FF):
            return "CJK"
        if 0xAC00 <= cp <= 0xD7AF:
            return "HANGUL"
        if 0x0600 <= cp <= 0x06FF or 0x0750 <= cp <= 0x077F:
            return "ARABIC"
        if 0x0370 <= cp <= 0x03FF:
            return "GREEK"
    if any(ord(ch) > 127 for ch in text):
        return "OTHER"
    return "LATIN"

def get_font(script, default_size=12):
    font_name = FALLBACK_FONTS.get(script, FALLBACK_FONTS["DEFAULT"])
    search_paths = [
        f"/usr/share/fonts/truetype/liberation/{font_name}",
        f"/usr/share/fonts/truetype/msttcorefonts/{font_name}",
        f"/usr/share/fonts/truetype/noto/{font_name}",
        f"/Library/Fonts/{font_name}",
        f"C:/Windows/Fonts/{font_name}",
        font_name
    ]
    for p in search_paths:
        try:
            if os.path.exists(p):
                return ImageFont.truetype(p, default_size)
        except Exception:
            pass
    # Fallbacks
    for fb in ["arialuni.ttf", "NotoSans-Regular.ttf"]:
        for p in search_paths:
            try:
                alt = p.replace(font_name, fb)
                if os.path.exists(alt):
                    return ImageFont.truetype(alt, default_size)
            except Exception:
                pass
    # Last resort
    try:
        return ImageFont.load_default()
    except Exception as e:
        raise RuntimeError(f"Could not load any font: {e}")

def wrap_text(line: str, font_size: int, img_width: int, margin: int) -> list:
    # empirical width per char estimate
    max_chars = max(1, int((img_width - 2 * margin) / (font_size * 0.6)))
    return textwrap.wrap(line, width=max_chars)

# =========================
# CSV Converters
# =========================

def csv_to_xls(csv_path, xls_path, chunksize=10000):
    try:
        ensure_parent_dir(xls_path)
        wb = Workbook(write_only=True)
        ws = wb.create_sheet("Sheet1")
        first = True
        for chunk in pd.read_csv(csv_path, chunksize=chunksize, dtype=str):
            if first:
                ws.append(list(chunk.columns))
                first = False
            for row in chunk.itertuples(index=False, name=None):
                ws.append(list(row))
        wb.save(xls_path)
        print(f"✅ CSV → XLSX: {xls_path}")
        return xls_path
    except Exception as e:
        print(f"❌ CSV to XLSX failed: {e}")
        return None

def csv_to_pdf(csv_path, pdf_path, margin=40, line_gap=14, font="Helvetica", size=10):
    try:
        ensure_parent_dir(pdf_path)
        c = canvas.Canvas(pdf_path, pagesize=A4)
        c.setFont(font, size)
        y = PAGE_H - margin
        max_chars = 180
        with open(csv_path, "r", encoding="utf-8", newline="") as f:
            reader = csv.reader(f)
            for row in reader:
                line = " | ".join([str(x) for x in row])
                chunks = [line[i:i+max_chars] for i in range(0, len(line), max_chars)] or [" "]
                for chunk in chunks:
                    if y < margin:
                        c.showPage(); c.setFont(font, size); y = PAGE_H - margin
                    c.drawString(margin, y, chunk)
                    y -= line_gap
        c.save()
        print(f"✅ CSV → PDF: {pdf_path}")
        return pdf_path
    except Exception as e:
        print(f"❌ CSV to PDF failed: {e}")
        return None

def csv_to_doc(csv_path, docx_path):
    try:
        ensure_parent_dir(docx_path)
        df = pd.read_csv(csv_path, dtype=str)
        doc = Document()
        table = doc.add_table(rows=1, cols=len(df.columns))
        for i, col in enumerate(df.columns):
            table.cell(0, i).text = col
        for row in df.itertuples(index=False, name=None):
            cells = table.add_row().cells
            for i, val in enumerate(row):
                cells[i].text = "" if pd.isna(val) else str(val)
        doc.save(docx_path)
        print(f"✅ CSV → DOCX: {docx_path}")
        return docx_path
    except Exception as e:
        print(f"❌ CSV to DOCX failed: {e}")
        return None

def csv_to_txt(csv_path, txt_path, delimiter="\t", chunksize=10000):
    try:
        ensure_parent_dir(txt_path)
        first = True
        with open(txt_path, "w", encoding="utf-8") as out:
            for chunk in pd.read_csv(csv_path, dtype=str, chunksize=chunksize):
                if first:
                    out.write(delimiter.join(chunk.columns) + "\n")
                    first = False
                for row in chunk.itertuples(index=False, name=None):
                    out.write(delimiter.join("" if pd.isna(x) else str(x) for x in row) + "\n")
        print(f"✅ CSV → TXT: {txt_path}")
        return txt_path
    except Exception as e:
        print(f"❌ CSV to TXT failed: {e}")
        return None

def csv_to_json(csv_path, json_path, chunksize=50000):
    try:
        ensure_parent_dir(json_path)
        first = True
        with open(json_path, "w", encoding="utf-8") as out:
            out.write("[")
            for chunk in pd.read_csv(csv_path, dtype=str, chunksize=chunksize):
                records = chunk.where(pd.notnull(chunk), None).to_dict(orient="records")
                for rec in records:
                    if not first: out.write(",\n")
                    json.dump(rec, out, ensure_ascii=False)
                    first = False
            out.write("]")
        print(f"✅ CSV → JSON: {json_path}")
        return json_path
    except Exception as e:
        print(f"❌ CSV to JSON failed: {e}")
        return None

# ===== CSV → Image (single or multi-image ZIP) with correct wrapping =====

def _render_chunk_to_image(lines_chunk, img_path, font_size=12, margin=40, line_height=18, img_width=1200, max_safe_height=30000):
    try:
        # Pre-measure wrapped line count
        total_wrapped_lines = 0
        wrapped_chunks = []
        for line in lines_chunk:
            script = detect_script_simple(line)
            font = get_font(script, font_size)
            wrapped = wrap_text(line, font_size, img_width, margin)
            if not wrapped:
                wrapped = [""]
            wrapped_chunks.append((wrapped, font))
            total_wrapped_lines += len(wrapped)

        img_height = margin * 2 + total_wrapped_lines * line_height
        if img_height > max_safe_height:
            print(f"⚠️ Image height {img_height}px is very large. Consider multi-image ZIP for better reliability.")

        img = Image.new("RGB", (img_width, img_height), "white")
        draw = ImageDraw.Draw(img)

        y = margin
        for wrapped, font in wrapped_chunks:
            for w_line in wrapped:
                draw.text((margin, y), w_line, fill="black", font=font)
                y += line_height

        img.save(img_path, "PNG")
        return img_path
    except Exception as e:
        print(f"❌ Render chunk failed: {e}")
        return None

def csv_to_image(csv_path, out_path, font_size=12, margin=40, line_height=18, img_width=1200, max_lines_per_img=60):
    """
    If CSV has > max_lines_per_img → ask:
      1. Single image
      2. Multi images (ZIP)
    """
    try:
        ensure_parent_dir(out_path)
        lines = []
        with open(csv_path, "r", encoding="utf-8", errors="ignore") as f:
            reader = csv.reader(f)
            for row in reader:
                lines.append(" | ".join(str(x) for x in row))
        if not lines:
            raise RuntimeError("CSV is empty.")

        total = len(lines)
        split = total > max_lines_per_img

        if split:
            print(f"CSV lines: {total} (> {max_lines_per_img})")
            print("Choose output option:")
            print("1. Single giant image (may be very tall)")
            print("2. Multiple images (zipped)")
            choice = input("Enter choice (1/2): ").strip()
            split = (choice != "1")

        images = []
        if not split:
            if not out_path.lower().endswith(".png"):
                out_path += ".png"
            path = _render_chunk_to_image(lines, out_path, font_size, margin, line_height, img_width)
            if path: images.append(path)
            print(f"✅ CSV → Image: {out_path}")
            return out_path
        else:
            chunks = [lines[i:i+max_lines_per_img] for i in range(0, total, max_lines_per_img)]
            img_dir = os.path.splitext(out_path)[0] + "_images"
            os.makedirs(img_dir, exist_ok=True)
            for i, chunk in enumerate(chunks):
                img_path = os.path.join(img_dir, f"page_{i+1}.png")
                _render_chunk_to_image(chunk, img_path, font_size, margin, line_height, img_width)
                images.append(img_path)
            zip_path = os.path.splitext(out_path)[0] + ".zip"
            with zipfile.ZipFile(zip_path, 'w') as zf:
                for img_file in images:
                    zf.write(img_file, os.path.basename(img_file))
            shutil.rmtree(img_dir)
            print(f"✅ CSV → Images ZIP: {zip_path}")
            return zip_path
    except Exception as e:
        print(f"❌ CSV to Image failed: {e}")
        return None

# =========================
# XLS Converters
# =========================

def xls_to_csv(xls_path, csv_path, sheet_name=None):
    try:
        ensure_parent_dir(csv_path)
        wb = load_workbook(xls_path, read_only=True, data_only=True)
        ws = wb[sheet_name] if sheet_name else wb.active
        with open(csv_path, "w", encoding="utf-8", newline="") as out:
            writer = csv.writer(out)
            for row in ws.iter_rows(values_only=True):
                writer.writerow(["" if v is None else v for v in row])
        print(f"✅ XLSX → CSV: {csv_path}")
        return csv_path
    except Exception as e:
        print(f"❌ XLSX to CSV failed: {e}")
        return None

def xls_to_doc(xls_path, docx_path, sheet_name=None):
    try:
        ensure_parent_dir(docx_path)
        wb = load_workbook(xls_path, read_only=True, data_only=True)
        ws = wb[sheet_name] if sheet_name else wb.active
        doc = Document()
        rows = list(ws.iter_rows(values_only=True))
        if not rows:
            doc.save(docx_path); print(f"⚠️ Sheet empty. DOCX saved: {docx_path}"); return docx_path
        table = doc.add_table(rows=1, cols=len(rows[0]))
        for i, val in enumerate(rows[0]):
            table.cell(0,i).text = "" if val is None else str(val)
        for r in rows[1:]:
            cells = table.add_row().cells
            for i, val in enumerate(r):
                cells[i].text = "" if val is None else str(val)
        doc.save(docx_path)
        print(f"✅ XLSX → DOCX: {docx_path}")
        return docx_path
    except Exception as e:
        print(f"❌ XLSX to DOCX failed: {e}")
        return None

def xls_to_txt(xls_path, txt_path, delimiter="\t", sheet_name=None):
    try:
        ensure_parent_dir(txt_path)
        wb = load_workbook(xls_path, read_only=True, data_only=True)
        ws = wb[sheet_name] if sheet_name else wb.active
        with open(txt_path, "w", encoding="utf-8") as out:
            for row in ws.iter_rows(values_only=True):
                out.write(delimiter.join("" if v is None else str(v) for v in row) + "\n")
        print(f"✅ XLSX → TXT: {txt_path}")
        return txt_path
    except Exception as e:
        print(f"❌ XLSX to TXT failed: {e}")
        return None

def xls_to_pdf(xls_path, pdf_path, margin=40, line_gap=14, font="Helvetica", size=10, sheet_name=None):
    try:
        ensure_parent_dir(pdf_path)
        wb = load_workbook(xls_path, read_only=True, data_only=True)
        ws = wb[sheet_name] if sheet_name else wb.active
        c = canvas.Canvas(pdf_path, pagesize=A4)
        c.setFont(font, size)
        y = PAGE_H - margin
        max_chars = 180
        for row in ws.iter_rows(values_only=True):
            line = " | ".join("" if v is None else str(v) for v in row)
            chunks = [line[i:i+max_chars] for i in range(0, len(line), max_chars)] or [" "]
            for ch in chunks:
                if y < margin:
                    c.showPage(); c.setFont(font, size); y = PAGE_H - margin
                c.drawString(margin, y, ch)
                y -= line_gap
        c.save()
        print(f"✅ XLSX → PDF: {pdf_path}")
        return pdf_path
    except Exception as e:
        print(f"❌ XLSX to PDF failed: {e}")
        return None

def xls_to_json(xls_path, json_path, sheet_name=None):
    try:
        ensure_parent_dir(json_path)
        wb = load_workbook(xls_path, read_only=True, data_only=True)
        ws = wb[sheet_name] if sheet_name else wb.active
        rows = list(ws.iter_rows(values_only=True))
        if not rows:
            with open(json_path, "w", encoding="utf-8") as out:
                out.write("[]")
            print(f"⚠️ Sheet empty. JSON saved: {json_path}")
            return json_path
        headers = [str(h) if h is not None else "" for h in rows[0]]
        data = []
        for r in rows[1:]:
            obj = {headers[i]: ("" if r[i] is None else r[i]) for i in range(len(headers))}
            data.append(obj)
        with open(json_path, "w", encoding="utf-8") as out:
            json.dump(data, out, ensure_ascii=False, indent=2)
        print(f"✅ XLSX → JSON: {json_path}")
        return json_path
    except Exception as e:
        print(f"❌ XLSX to JSON failed: {e}")
        return None

# =========================
# TXT Converters
# =========================

def txt_to_pdf(txt_path, pdf_path, margin=40, line_gap=14, font="Helvetica", size=10):
    try:
        ensure_parent_dir(pdf_path)
        c = canvas.Canvas(pdf_path, pagesize=A4)
        c.setFont(font, size)
        y = PAGE_H - margin
        max_chars = 180
        with open(txt_path, "r", encoding="utf-8", errors="ignore") as f:
            for raw in f.read().splitlines():
                line = raw or " "
                for chunk in [line[i:i+max_chars] for i in range(0, len(line), max_chars)] or [" "]:
                    if y < margin:
                        c.showPage(); c.setFont(font, size); y = PAGE_H - margin
                    c.drawString(margin, y, chunk)
                    y -= line_gap
        c.save()
        print(f"✅ TXT → PDF: {pdf_path}")
        return pdf_path
    except Exception as e:
        print(f"❌ TXT to PDF failed: {e}")
        return None

def txt_to_doc(txt_path, docx_path):
    try:
        ensure_parent_dir(docx_path)
        doc = Document()
        with open(txt_path, "r", encoding="utf-8", errors="ignore") as f:
            for line in f.read().splitlines():
                doc.add_paragraph(line)
        doc.save(docx_path)
        print(f"✅ TXT → DOCX: {docx_path}")
        return docx_path
    except Exception as e:
        print(f"❌ TXT to DOCX failed: {e}")
        return None

def txt_to_json(txt_path, json_path):
    try:
        ensure_parent_dir(json_path)
        with open(txt_path, "r", encoding="utf-8", errors="ignore") as f:
            lines = [ln.rstrip("\n") for ln in f.readlines()]
        with open(json_path, "w", encoding="utf-8") as out:
            json.dump(lines, out, ensure_ascii=False, indent=2)
        print(f"✅ TXT → JSON: {json_path}")
        return json_path
    except Exception as e:
        print(f"❌ TXT to JSON failed: {e}")
        return None

def txt_to_csv(txt_path, csv_path, delimiter="\t"):
    try:
        ensure_parent_dir(csv_path)
        with open(txt_path, "r", encoding="utf-8", errors="ignore") as f, \
             open(csv_path, "w", encoding="utf-8", newline="") as out:
            writer = csv.writer(out)
            for line in f.read().splitlines():
                parts = line.split(delimiter) if delimiter else [line]
                writer.writerow(parts)
        print(f"✅ TXT → CSV: {csv_path}")
        return csv_path
    except Exception as e:
        print(f"❌ TXT to CSV failed: {e}")
        return None

def txt_to_image(txt_path, out_path, font_size=12, margin=40, line_height=18, img_width=1200, max_lines_per_img=60):
    try:
        ensure_parent_dir(out_path)
        with open(txt_path, "r", encoding="utf-8", errors="ignore") as f:
            lines = [ln.rstrip("\n") for ln in f.readlines()]
        if not lines:
            raise RuntimeError("TXT is empty.")

        total = len(lines)
        split = total > max_lines_per_img
        if split:
            print(f"TXT lines: {total} (> {max_lines_per_img})")
            print("Choose output option:")
            print("1. Single giant image")
            print("2. Multiple images (zipped)")
            choice = input("Enter choice (1/2): ").strip()
            split = (choice != "1")

        if not split:
            if not out_path.lower().endswith(".png"):
                out_path += ".png"
            _render_chunk_to_image(lines, out_path, font_size, margin, line_height, img_width)
            print(f"✅ TXT → Image: {out_path}")
            return out_path
        else:
            chunks = [lines[i:i+max_lines_per_img] for i in range(0, total, max_lines_per_img)]
            img_dir = os.path.splitext(out_path)[0] + "_images"
            os.makedirs(img_dir, exist_ok=True)
            images = []
            for i, ch in enumerate(chunks):
                img_path = os.path.join(img_dir, f"page_{i+1}.png")
                _render_chunk_to_image(ch, img_path, font_size, margin, line_height, img_width)
                images.append(img_path)
            zip_path = os.path.splitext(out_path)[0] + ".zip"
            with zipfile.ZipFile(zip_path, 'w') as zf:
                for img_file in images:
                    zf.write(img_file, os.path.basename(img_file))
            shutil.rmtree(img_dir)
            print(f"✅ TXT → Images ZIP: {zip_path}")
            return zip_path
    except Exception as e:
        print(f"❌ TXT to Image failed: {e}")
        return None

# =========================
# JSON Converters
# =========================

def json_to_csv(json_path, csv_path):
    try:
        ensure_parent_dir(csv_path)
        with open(json_path, "r", encoding="utf-8") as f:
            data = json.load(f)
        if isinstance(data, dict):
            data = [data]
        df = pd.DataFrame(data)
        df.to_csv(csv_path, index=False)
        print(f"✅ JSON → CSV: {csv_path}")
        return csv_path
    except Exception as e:
        print(f"❌ JSON to CSV failed: {e}")
        return None

def json_to_xls(json_path, xls_path):
    try:
        ensure_parent_dir(xls_path)
        with open(json_path, "r", encoding="utf-8") as f:
            data = json.load(f)
        if isinstance(data, dict):
            data = [data]
        df = pd.DataFrame(data)
        df.to_excel(xls_path, index=False)
        print(f"✅ JSON → XLSX: {xls_path}")
        return xls_path
    except Exception as e:
        print(f"❌ JSON to XLSX failed: {e}")
        return None

def json_to_txt(json_path, txt_path):
    try:
        ensure_parent_dir(txt_path)
        with open(json_path, "r", encoding="utf-8") as f:
            data = json.load(f)
        out_txt = json.dumps(data, ensure_ascii=False, indent=2)
        with open(txt_path, "w", encoding="utf-8") as out:
            out.write(out_txt)
        print(f"✅ JSON → TXT: {txt_path}")
        return txt_path
    except Exception as e:
        print(f"❌ JSON to TXT failed: {e}")
        return None

# =========================
# PDF & DOCX Helpers
# =========================

def pdf_to_txt(pdf_path, txt_path):
    try:
        ensure_parent_dir(txt_path)
        text = pdf_extract_text(pdf_path) or ""
        with open(txt_path, "w", encoding="utf-8") as out:
            out.write(text)
        print(f"✅ PDF → TXT: {txt_path}")
        return txt_path
    except Exception as e:
        print(f"❌ PDF to TXT failed: {e}")
        return None

def doc_to_txt(docx_path, txt_path):
    try:
        ensure_parent_dir(txt_path)
        doc = Document(docx_path)
        lines = []
        for p in doc.paragraphs:
            lines.append(p.text)
        with open(txt_path, "w", encoding="utf-8") as out:
            out.write("\n".join(lines))
        print(f"✅ DOCX → TXT: {txt_path}")
        return txt_path
    except Exception as e:
        print(f"❌ DOCX to TXT failed: {e}")
        return None

def doc_to_pdf(docx_path, pdf_path, margin=40, line_gap=14, font="Helvetica", size=10):
    """Simple text-based DOCX → PDF (no layout preservation)."""
    try:
        ensure_parent_dir(pdf_path)
        tmp_txt = os.path.splitext(pdf_path)[0] + ".__tmp_doc.txt"
        if not doc_to_txt(docx_path, tmp_txt):
            return None
        res = txt_to_pdf(tmp_txt, pdf_path, margin=margin, line_gap=line_gap, font=font, size=size)
        try:
            if os.path.exists(tmp_txt):
                os.remove(tmp_txt)
        except Exception:
            pass
        if res:
            print(f"✅ DOCX → PDF: {pdf_path}")
        return res
    except Exception as e:
        print(f"❌ DOCX to PDF failed: {e}")
        return None

# =========================
# Translation
# =========================

def translate_text(text, src_lang="auto", dest_lang="en"):
    translated_chunks = []
    chunk_size = 4000  # Google limit-ish
    for i in range(0, len(text), chunk_size):
        chunk = text[i:i+chunk_size]
        translated = GoogleTranslator(source=src_lang, target=dest_lang).translate(chunk)
        translated_chunks.append(translated)
    return "\n".join(translated_chunks)

def translate_file(file_path, out_path, src_lang="auto", dest_lang="en"):
    try:
        ensure_parent_dir(out_path)
        ext = os.path.splitext(file_path)[1].lower()
        if ext == ".pdf":
            temp_txt = os.path.splitext(out_path)[0] + ".__tmp_in.txt"
            txt_extracted = pdf_to_txt(file_path, temp_txt)
            if not txt_extracted: return None
            with open(temp_txt, "r", encoding="utf-8") as f:
                content = f.read()
            try:
                os.remove(temp_txt)
            except Exception:
                pass
        elif ext == ".docx":
            temp_txt = os.path.splitext(out_path)[0] + ".__tmp_in.txt"
            txt_extracted = doc_to_txt(file_path, temp_txt)
            if not txt_extracted: return None
            with open(temp_txt, "r", encoding="utf-8") as f:
                content = f.read()
            try:
                os.remove(temp_txt)
            except Exception:
                pass
        else:
            with open(file_path, "r", encoding="utf-8") as f:
                content = f.read()

        translated_text = translate_text(content, src_lang, dest_lang)
        with open(out_path, "w", encoding="utf-8") as f:
            f.write(translated_text)
        print(f"✅ Translated → {out_path}")
        return out_path
    except Exception as e:
        print(f"❌ Translation failed: {e}")
        return None

# =========================
# Dispatcher
# =========================

CONVERTERS = {
    # CSV →
    ("csv", "xlsx"): csv_to_xls,
    ("csv", "pdf"): csv_to_pdf,
    ("csv", "docx"): csv_to_doc,
    ("csv", "txt"): csv_to_txt,
    ("csv", "json"): csv_to_json,
    ("csv", "png"): csv_to_image,

    # XLSX →
    ("xlsx", "csv"): xls_to_csv,
    ("xlsx", "docx"): xls_to_doc,
    ("xlsx", "txt"): xls_to_txt,
    ("xlsx", "pdf"): xls_to_pdf,
    ("xlsx", "json"): xls_to_json,

    # TXT →
    ("txt", "pdf"): txt_to_pdf,
    ("txt", "docx"): txt_to_doc,
    ("txt", "json"): txt_to_json,
    ("txt", "csv"): txt_to_csv,
    ("txt", "png"): txt_to_image,

    # JSON →
    ("json", "csv"): json_to_csv,
    ("json", "xlsx"): json_to_xls,
    ("json", "txt"): json_to_txt,

    # DOCX → (helpers)
    ("docx", "txt"): doc_to_txt,
    ("docx", "pdf"): doc_to_pdf,

    # PDF → (helpers)
    ("pdf", "txt"): pdf_to_txt,
}

def infer_ext(path):
    return os.path.splitext(path)[1].lower().strip(".") if "." in os.path.basename(path) else ""

def run_conversion(src_fmt, dst_fmt, in_path, out_path):
    key = (src_fmt.lower(), dst_fmt.lower())
    if key not in CONVERTERS:
        print(f"❌ Conversion not supported: {src_fmt} → {dst_fmt}")
        return None
    return CONVERTERS[key](in_path, out_path)

# =========================
# CLI Menu
# =========================

def cli_menu():
    while True:
        print("\n===== File Toolkit =====")
        print("1. Convert File")
        print("2. Translate File")
        print("3. Exit")
        choice = input("Enter choice: ").strip()

        if choice == "1":
            try:
                in_path = input("Enter input file path: ").strip()
                if not os.path.exists(in_path):
                    print("❌ Input file not found.")
                    continue

                src_fmt = input("Enter source format (auto to detect): ").strip().lower()
                if src_fmt in ("", "auto"):
                    src_fmt = infer_ext(in_path)
                    if not src_fmt:
                        print("❌ Could not infer source format. Provide explicitly.")
                        continue

                print("Supported targets for your source may include: csv, xlsx, pdf, docx, txt, json, png")
                dst_fmt = input("Enter target format: ").strip().lower()
                if not dst_fmt:
                    print("❌ Target format required.")
                    continue

                out_path = input("Enter output file path (leave blank to auto): ").strip()
                if not out_path:
                    base, _ = os.path.splitext(in_path)
                    out_path = f"{base}.{dst_fmt}"

                result = run_conversion(src_fmt, dst_fmt, in_path, out_path)
                if result:
                    print(f"✅ Converted successfully: {result}")
                    # Optional auto-delete (uncomment if needed)
                    # schedule_delete(result)
                else:
                    print("❌ Conversion failed.")
            except Exception as e:
                print(f"❌ Error: {e}")

        elif choice == "2":
            try:
                in_path = input("Enter file path to translate: ").strip()
                if not os.path.exists(in_path):
                    print("❌ Input file not found.")
                    continue
                src_lang = input("Source language (default 'auto'): ").strip() or "auto"
                dst_lang = input("Target language (default 'en'): ").strip() or "en"
                out_path = input("Output text file path (e.g., translated.txt): ").strip()
                if not out_path:
                    base, _ = os.path.splitext(in_path)
                    out_path = f"{base}.translated.txt"

                result = translate_file(in_path, out_path, src_lang=src_lang, dest_lang=dst_lang)
                if result:
                    print(f"✅ Translated successfully: {result}")
                    # schedule_delete(result)
                else:
                    print("❌ Translation failed.")
            except Exception as e:
                print(f"❌ Error: {e}")

        elif choice == "3":
            print("👋 Exiting...")
            break
        else:
            print("❌ Invalid choice!")

# =========================
# Entry
# =========================

if __name__ == "__main__":
    # If you also want a non-interactive mode, you can add argparse here.
    # For now, we use the requested menu UI:
    cli_menu()
