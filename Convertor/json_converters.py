# json_converters.py
from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph
from reportlab.lib.pagesizes import A4
from reportlab.lib import colors
from pdf2image import convert_from_path
from tqdm import tqdm   # pip install tqdm
from reportlab.lib.styles import getSampleStyleSheet
from docx import Document
from PIL import Image, ImageDraw, ImageFont
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
import threading
import json, csv
import pandas as pd
from openpyxl import Workbook
import json
from reportlab.pdfgen import canvas
from docx.shared import Pt
from docx.oxml.ns import qn
import os, zipfile, textwrap, uuid
import tempfile
import shutil
import ijson
import pdfkit
import pypandoc
from weasyprint import HTML

def _ensure_dir(file_path):
    os.makedirs(os.path.dirname(file_path), exist_ok=True)

def register_unicode_font(font_files=None):
    """
    Registers first available TrueType font for Unicode support.
    Returns the font name to be used in PDF.
    """
    if font_files is None:
        font_files = ["NotoSans-Regular.ttf", "NotoSansDevanagari.ttf"]
    for f in font_files:
        if os.path.exists(f):
            try:
                pdfmetrics.registerFont(TTFont("UnicodeFont", f))
                return "UnicodeFont"
            except Exception as e:
                print(f"⚠️ Could not register font {f}: {e}")
    print("⚠️ No Unicode font found, using default Helvetica")
    return "Helvetica"

#------------------------- Json To PDF -----------------------------------------------#

def json_to_pdf(json_path, output_pdf):
    try:

        # --- Only allow .json files ---
        if not json_path.lower().endswith(".json"):
            print(f"❌ Unsupported file type: {json_path}. Only .json files are supported.")

        if not os.path.exists(json_path):
            raise RuntimeError("Input file not found: " + json_path)

        # Detect JSON type: normal vs JSON lines
        data_list = []
        with open(json_path, "r", encoding="utf-8") as f:
            first_char = f.read(1)
            f.seek(0)
            if first_char == "{":
                # Could be JSON Lines or single JSON object
                try:
                    # Try loading whole file as a JSON object
                    data = json.load(f)
                    if isinstance(data, dict):
                        data_list.append(data)
                    elif isinstance(data, list):
                        data_list = data
                except json.JSONDecodeError:
                    # Treat as JSON Lines
                    f.seek(0)
                    for line in f:
                        line = line.strip()
                        if line:
                            data_list.append(json.loads(line))
            elif first_char == "[":
                data_list = json.load(f)
            else:
                raise RuntimeError("Invalid JSON format")

        # Convert to pretty JSON string
        json_text = json.dumps(data_list, indent=4, ensure_ascii=False)

        # HTML content for PDF
        html_content = f"""
        <html>
        <head>
            <meta charset="utf-8">
            <style>
                body {{ font-family: 'Noto Sans', 'Devanagari Sans', monospace; font-size:12pt; }}
                pre {{ white-space: pre-wrap; word-wrap: break-word; }}
            </style>
        </head>
        <body>
            <pre>{json_text}</pre>
        </body>
        </html>
        """

        tmp_html = tempfile.mktemp(suffix=".html")
        with open(tmp_html, "w", encoding="utf-8") as f:
            f.write(html_content)

        pdfkit.from_file(tmp_html, output_pdf)
        os.remove(tmp_html)
        return output_pdf

    except Exception as e:
      print(f"Error{e}")

#------------------------- Json To DOCX ----------------------------------------#
def json_to_doc(json_file, output_file):
    try:
        if not json_file.endswith(".json"):
            raise RuntimeError("❌ Only JSON files are supported!")

        if not os.path.exists(json_file):
            raise RuntimeError("❌ File not found!")

        data_list = []
        with open(json_file, "r", encoding="utf-8") as f:
            first_char = f.read(1)
            f.seek(0)
            if first_char == "{":
                try:
                    # Normal JSON
                    data = json.load(f)
                    if isinstance(data, dict):
                        data_list.append(data)
                    elif isinstance(data, list):
                        data_list = data
                except json.JSONDecodeError:
                    # NDJSON
                    f.seek(0)
                    for line in f:
                        line = line.strip()
                        if line:
                            data_list.append(json.loads(line))
            elif first_char == "[":
                data_list = json.load(f)
            else:
                raise RuntimeError("❌ Invalid JSON format")

        if not data_list:
            raise RuntimeError("❌ No data found in JSON.")

        # --- Create DOCX ---
        doc = Document()
        style = doc.styles['Normal']
        font = style.font
        font.name = 'Nirmala UI'   # Hindi-friendly font (fallback)
        font.size = Pt(11)
        style._element.rPr.rFonts.set(qn('w:eastAsia'), 'Nirmala UI')

        doc.add_heading("JSON Data Export", level=1)

        # --- Output formatting ---
        if isinstance(data_list, list) and all(isinstance(item, dict) for item in data_list):
            # Multiple dicts → print each as formatted block
            for idx, obj in enumerate(data_list, 1):
                doc.add_paragraph(f"Record {idx}:", style='Heading 2')
                pretty = json.dumps(obj, indent=4, ensure_ascii=False)
                doc.add_paragraph(pretty)
        else:
            # Single object or nested → dump as pretty JSON
            pretty = json.dumps(data_list, indent=4, ensure_ascii=False)
            doc.add_paragraph(pretty)

        # --- Save DOCX ---
        doc.save(output_file)
        return output_file

    except Exception as e:
        raise RuntimeError(f"❌ Conversion failed: {e}")


# ------------------ Json TO Image ------------------------------------------
def is_ndjson(json_path):
    """Detect NDJSON (multiple lines with valid JSON objects)"""
    with open(json_path, "r", encoding="utf-8") as f:
        try:
            json.load(f)   # normal JSON
            return False
        except json.JSONDecodeError:
            return True

def json_to_image(json_path, output_path):
    try:
        if not json_path.endswith(".json"):
            raise RuntimeError("❌ Only JSON files are supported!")
            
        if not os.path.exists(json_path):
            raise RuntimeError("❌ File not found!")

        ndjson_mode = is_ndjson(json_path)

        print("Choose conversion mode:")
        print("1. Single big image (whole JSON in one PNG)")
        print("2. Split into 30-line chunks (ZIP of PNGs)")
        choice = input("Enter choice (1/2): ").strip()

        if ndjson_mode:
            print("📂 NDJSON detected → routing via json_to_pdf()...")
            tmp_pdf = tempfile.mktemp(suffix=".pdf")
            json_to_pdf(json_path, tmp_pdf)   # use your fast code

            # convert PDF → images
            images = convert_from_path(tmp_pdf, dpi=100)

            if choice == "1":
                # merge pages into one long PNG
                total_height = sum(i.height for i in images)
                max_width = max(i.width for i in images)
                big_img = Image.new("RGB", (max_width, total_height), "white")
                y_offset = 0
                for img in images:
                    big_img.paste(img, (0, y_offset))
                    y_offset += img.height
                big_img.save(output_path, "PNG")
                print(f"✅ NDJSON → single long image: {output_path}")
                return output_path
            else:
                # each page separate → zip
                tmpdir = tempfile.mkdtemp()
                img_files = []
                base, ext = os.path.splitext(output_path)
                for i, page in enumerate(images, 1):
                    out_file = os.path.join(tmpdir, f"page_{i}.png")
                    page.save(out_file, "PNG")
                    img_files.append(out_file)

                zip_path = output_path.replace(".png", ".zip")
                with zipfile.ZipFile(zip_path, "w") as zipf:
                    for img in img_files:
                        zipf.write(img, os.path.basename(img))

                print(f"✅ NDJSON → images zipped: {zip_path}")
                return zip_path

        else:
            # ✅ Normal JSON → existing HTML flow
            with open(json_path, "r", encoding="utf-8") as f:
                data_json = json.load(f)
            if isinstance(data_json, dict):
                data_json = [data_json]
            pretty_json = json.dumps(data_json, indent=4, ensure_ascii=False)

            if choice == "1":
                return save_html_as_image(pretty_json, output_path, long_mode=True)
            else:
                lines = pretty_json.split("\n")
                tmpdir = tempfile.mkdtemp()
                img_files = []
                chunk_size = 30
                for i in range(0, len(lines), chunk_size):
                    chunk = "\n".join(lines[i:i + chunk_size])
                    img_file = os.path.join(tmpdir, f"chunk_{i//chunk_size + 1}.png")
                    save_html_as_image(chunk, img_file, long_mode=False)
                    img_files.append(img_file)

                zip_path = output_path.replace(".png", ".zip")
                with zipfile.ZipFile(zip_path, "w") as zipf:
                    for img in img_files:
                        zipf.write(img, os.path.basename(img))

                print(f"✅ Converted successfully (chunks zipped): {zip_path}")
                return zip_path

    except Exception as e:
        raise RuntimeError(f"❌ Conversion failed: {e}")

def save_html_as_image(text, out_path, long_mode=False):
    html_content = f"""
    <html>
    <head>
        <meta charset="UTF-8">
        <style>
            body {{
                font-family: 'Nirmala UI','Noto Sans Devanagari','Noto Sans','Arial Unicode MS',sans-serif;
                font-size: 14pt;
                white-space: pre-wrap;
            }}
            pre {{
                white-space: pre-wrap;
                word-wrap: break-word;
            }}
        </style>
    </head>
    <body><pre>{text}</pre></body></html>
    """
    tmp_html = tempfile.mktemp(suffix=".html")
    with open(tmp_html, "w", encoding="utf-8") as f:
        f.write(html_content)

    tmp_pdf = out_path.replace(".png", f"_{uuid.uuid4().hex}.pdf")
    HTML(tmp_html).write_pdf(tmp_pdf)
    images = convert_from_path(tmp_pdf, dpi=100)

    if long_mode and len(images) > 1:
        total_height = sum(i.height for i in images)
        max_width = max(i.width for i in images)
        big_img = Image.new("RGB", (max_width, total_height), "white")
        y_offset = 0
        for img in images:
            big_img.paste(img, (0, y_offset))
            y_offset += img.height
        big_img.save(out_path, "PNG")
        print(f"✅ Saved single long image: {out_path}")
        return out_path
    else:
        images[0].save(out_path, "PNG")
        print(f"✅ Saved image: {out_path}")
        return out_path

#----------------------- Json To TXT-------------------------------#
def json_to_txt(json_path, txt_path):
    """
    Convert JSON/NDJSON to TXT.
    - Handles large files (50MB+).
    - Simple JSON => pretty indented text.
    - NDJSON => record-by-record streaming.
    - UTF-8 safe (Hindi + English).
    """
    try :

        if not json_path.endswith(".json"):
              raise RuntimeError("❌ Only JSON files are supported!")

        if not os.path.exists(json_path):
              raise RuntimeError("❌ File not found!")

        with open(json_path, "r", encoding="utf-8") as f, \
            open(txt_path, "w", encoding="utf-8") as out:

            try:
                # Try loading as a whole JSON (simple JSON)
                data = json.load(f)
                # pretty print
                pretty = json.dumps(data, ensure_ascii=False, indent=4)
                out.write(pretty + "\n")

            except json.JSONDecodeError:
                # Fallback: NDJSON mode
                f.seek(0)
                line_no = 0
                for line in f:
                    line_no += 1
                    line = line.strip()
                    if not line:
                        continue
                    try:
                        obj = json.loads(line)
                        # pretty-print each object with indentation
                        pretty = json.dumps(obj, ensure_ascii=False, indent=4)
                        out.write(f"--- Record {line_no} ---\n{pretty}\n\n")
                    except Exception:
                        out.write(f"⚠️ Skipping bad line {line_no}: {line[:50]}\n")

        return txt_path

    except Exception as e:
        print(f"❌ JSON to TXT failed: {e}")
        return None
        
#-------------------- Json To CSV ----------------------------------#
def json_to_csv(json_path, csv_path):
    """
    Convert JSON or NDJSON into CSV.
    - Handles large files (50MB+).
    - UTF-8 safe (Hindi + English).
    - Supports list-of-objects JSON & NDJSON.
    """
    try:
        if not json_path.endswith(".json"):
            raise RuntimeError("❌ Only JSON files are supported!")

        if not os.path.exists(json_path):
            raise RuntimeError("❌ File not found!")

        with open(csv_path, "w", encoding="utf-8-sig", newline="") as out:
            # utf-8-sig ensures Excel also reads Hindi properly
            writer = None

            try:
                # ---- Try as normal JSON (array of objects) ----
                with open(json_path, "r", encoding="utf-8") as f:
                    data = json.load(f)

                if isinstance(data, dict):
                    # if root is dict → wrap into list
                    data = [data]

                # Ensure it's list of dicts
                if not isinstance(data, list):
                    raise ValueError("JSON is not a list-of-objects")

                # Write CSV
                for idx, obj in enumerate(data):
                    if writer is None:
                        writer = csv.DictWriter(out, fieldnames=list(obj.keys()))
                        writer.writeheader()
                    writer.writerow(obj)

            except json.JSONDecodeError:
                # ---- NDJSON fallback (streaming) ----
                with open(json_path, "r", encoding="utf-8") as f:
                    for obj in ijson.items(f, "", multiple_values=True):
                        if isinstance(obj, dict):
                            if writer is None:
                                writer = csv.DictWriter(out, fieldnames=list(obj.keys()))
                                writer.writeheader()
                            writer.writerow(obj)
                        else:
                            # if NDJSON line is not dict
                            if writer is None:
                                writer = csv.writer(out)
                                writer.writerow(["value"])
                            writer.writerow([obj])

        return csv_path

    except Exception as e:
        print(f"❌ JSON to CSV failed: {e}")
        return None
        
# ----------------------- JSON To XLS --------------------------------------#
def flatten_json(obj, parent_key="", sep="."):
    """Flatten nested JSON (dicts/lists) into key-value pairs."""
    items = []
    if isinstance(obj, dict):
        for k, v in obj.items():
            new_key = f"{parent_key}{sep}{k}" if parent_key else k
            items.extend(flatten_json(v, new_key, sep=sep).items())
    elif isinstance(obj, list):
        for i, v in enumerate(obj):
            new_key = f"{parent_key}[{i}]"
            items.extend(flatten_json(v, new_key, sep=sep).items())
    else:
        items.append((parent_key, obj))
    return dict(items)

def json_to_xls(json_path, xls_path):
    """
    Convert JSON or NDJSON into XLSX.
    - Handles large files via streaming
    - UTF-8 safe (Hindi/English)
    - Flattens nested JSON
    """
    try:
        if not json_path.endswith(".json"):
            raise RuntimeError("❌ Only JSON files are supported!")

        if not os.path.exists(json_path):
            raise RuntimeError("❌ File not found!")

        wb = Workbook(write_only=True)
        ws = wb.create_sheet("Sheet1")

        headers_written = False
        headers = None

        with open(json_path, "r", encoding="utf-8") as f:
            try:
                # ---- simple JSON ----
                data = json.load(f)
                if isinstance(data, dict):
                    data = [data]
                if not isinstance(data, list):
                    raise ValueError("JSON is not a list of objects.")

                for obj in data:
                    if not isinstance(obj, dict):
                        continue
                    flat = flatten_json(obj)
                    if not headers_written:
                        headers = list(flat.keys())
                        ws.append(headers)
                        headers_written = True
                    ws.append([flat.get(h, "") for h in headers])

            except json.JSONDecodeError:
                # ---- NDJSON ----
                f.seek(0)
                for obj in ijson.items(f, "", multiple_values=True):
                    if not isinstance(obj, dict):
                        continue
                    flat = flatten_json(obj)
                    if not headers_written:
                        headers = list(flat.keys())
                        ws.append(headers)
                        headers_written = True
                    ws.append([flat.get(h, "") for h in headers])

        wb.save(xls_path)
        wb.close()

        if not os.path.exists(xls_path):
            return None
        return xls_path

    except Exception as e:
        print(f"❌ JSON to XLSX failed: {e}")
        return None

def schedule_delete(file_path, delay=300):  # 5 minutes
    def delete_file():
        if os.path.exists(file_path):
            try:
                if os.path.isdir(file_path):
                    import shutil
                    shutil.rmtree(file_path)
                else:
                    os.remove(file_path)
                print(f"🗑 Deleted: {file_path}")
            except Exception as e:
                print(f"❌ Failed to delete {file_path}: {e}")
    threading.Timer(delay, delete_file).start()

# ---------------- CLI MENU ---------------- #

def main():
    try:
        from google.colab import files
    except ImportError:
        files = None

    while True:
        print("\n==== File Converter =====")
        print("1. Convert JSON File")
        print("2. Exit")
        choice = input("Enter choice: ").strip()

        if choice == "1":

            # Upload file (Colab) or enter path (offline)
            if files:
                uploaded = files.upload()
                pdf_file = list(uploaded.keys())[0]
                #json_file = input("Enter path of JSON file: ").strip()
            else:
                json_file = input("Enter path of DOC/DOCX file: ").strip()

            print("\nSupported conversions: ")
            print("1. JSON → PDF")
            print("2. JSON → DOCX")
            print("3. JSON → PNG (image)")
            print("4. JSON → TXT")
            print("5. JSON → CSV")
            print("6. JSON → XLSX")
            fmt_choice = input("Select target format (1-6): ").strip()

            base, _ = os.path.splitext(json_file)
            out_file = None

            try:
                if fmt_choice == "1":
                    out_file = base + ".pdf"
                    json_to_pdf(json_file, out_file)
                elif fmt_choice == "2":
                    out_file = base + ".docx"
                    json_to_doc(json_file, out_file)
                elif fmt_choice == "3":
                    out_file = base + ".png"
                    json_to_image(json_file, out_file)
                elif fmt_choice == "4":
                    out_file = base + ".txt"
                    json_to_txt(json_file, out_file)
                elif fmt_choice == "5":
                    out_file = base + ".csv"
                    json_to_csv(json_file, out_file)
                elif fmt_choice == "6":
                    out_file = base + ".xlsx"
                    json_to_xls(json_file, out_file)
                else:
                    print("❌ Invalid choice!")
                    continue

                print(f"✅ Converted successfully: {out_file}")

                if out_file:
                    schedule_delete(out_file)

                if files and out_file:
                    files.download(out_file)

            except Exception as e:
                print("❌ Conversion failed:", e)

        elif choice == "2":
            print("👋 Exiting...")
            break
        else:
            print("❌ Invalid choice!")

if __name__ == "__main__":
    main()
